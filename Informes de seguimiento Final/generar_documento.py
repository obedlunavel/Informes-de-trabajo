from docx import Document
from docx.shared import Inches
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor  # Asegúrate de importar RGBColor
from docx import Document
import sys
import os

def resource_path(relative_path):
    try:
    # PyInstaller usa este directorio temporal
        base_path = sys._MEIPASS
    except AttributeError:
    # Ruta normal si no está empaquetado
        base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)
    
def generar_documento_word(datos, nombre_archivo):
    print("Generando documento Word...")  # Mensaje de depuración
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Función para eliminar los bordes de una celda específica
    def remove_cell_borders(cell, borders_to_remove):
        """
        Elimina los bordes especificados de una celda.

        :param cell: La celda a la que se le eliminarán los bordes.
        :param borders_to_remove: Una lista de los bordes a eliminar. Ej: ['top', 'left', 'bottom', 'right']
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # Crear el elemento w:tcBorders si no existe
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # Eliminar los bordes especificados
        for border in borders_to_remove:
            border_element = tcBorders.find(qn(f'w:{border}'))
            if border_element is None:
                border_element = OxmlElement(f'w:{border}')
                tcBorders.append(border_element)
            border_element.set(qn('w:val'), 'nil')  # Eliminar el borde

    # Función para alinear el contenido de la celda verticalmente al centro
    def set_vertical_alignment(cell, align="center"):
        """
        Alinea el contenido de la celda verticalmente.

        :param cell: La celda a la que se le aplicará la alineación vertical.
        :param align: La alineación vertical ('top', 'center', 'bottom').
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), align)
        tcPr.append(tcVAlign)

    # --- Encabezado ---
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # --- Añadir imágenes al header ---
    section = doc.sections[0]
    header = section.header

    # Crear una tabla en el encabezado con 1 fila y 2 columnas
    table_header = header.add_table(rows=1, cols=2, width=Inches(6))
    table_header.autofit = True
    # Ajustar el alto de las filas en el encabezado
    for row in table_header.rows:
        row.height = Inches(0.1)  # Ajusta el alto de las filas

    # Ajustar el ancho de las columnas
    for idx, col in enumerate(table_header.columns):
        col.width = Inches(5)  # Cada columna tendrá un ancho de 3 pulgadas

    # Insertar la primera imagen en la celda izquierda
    image_path_left = resource_path("Cabezerai.png")
    cell_left = table_header.cell(0, 0)
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_left.paragraphs[0].add_run().add_picture(image_path_left, width=Inches(1.6), height=Inches(0.7))  # Ajusta el alto de la imagen

    # Insertar la segunda imagen en la celda derecha
    image_path_right = resource_path("Cabezerai2.png")
    cell_right = table_header.cell(0, 1)
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_right.paragraphs[0].add_run().add_picture(image_path_right, width=Inches(1.2), height=Inches(0.7))  # Ajusta el alto de la imagen


    # Ajustar el alto de las filas
    for row in table.rows:
        row.height = Inches(0.3)  # Establece el alto de las filas a 0.5 pulgadas

    # Ajustar el ancho de las columnas
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(3)  # Establece el ancho de las columnas a 3 pulgadas

    # Eliminar bordes específicos
    remove_cell_borders(table.cell(0, 0), ['left', 'top'])
    remove_cell_borders(table.cell(1, 0), ['left', 'top', 'bottom'])
    remove_cell_borders(table.cell(2, 0), ['left', 'bottom'])

    # Fusionar celdas para el logo (3x1)
    cell_logo = table.cell(0, 0)
    cell_logo.merge(table.cell(2, 0))
    set_vertical_alignment(cell_logo, "center")
    cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    image_path = resource_path("Tabla.png")
    cell_logo.paragraphs[0].add_run().add_picture(image_path, width=Inches(2.5))

    # Dividir la celda del lado derecho en 3 celdas verticales
    cell_fecha = table.cell(0, 1)  # Celda para la fecha
    cell_fecha.text = f'FECHA: {datos["fecha"]}'
    cell_fecha.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_vertical_alignment(cell_fecha, "center")

    cell_unidad = table.cell(1, 1)  # Celda para la unidad
    cell_unidad.text = f'UNIDAD: {datos["unidad"]}'
    cell_unidad.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_vertical_alignment(cell_unidad, "center")

    cell_expediente = table.cell(2, 1)  # Celda para el expediente
    cell_expediente.text = f'NO. EXPEDIENTE: {datos["expediente"]}'
    cell_expediente.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_vertical_alignment(cell_expediente, "center")

    # Fusionar celdas para el título "INFORME DE SEGUIMIENTO"
    cell_titulo = table.cell(3, 0)
    cell_titulo.merge(table.cell(3, 1))
    cell_titulo.text = 'INFORME DE SEGUIMIENTO'
    # Obtener el primer párrafo de la celda
    paragraph = cell_titulo.paragraphs[0]

    # Alinear el texto al centro
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Poner el texto en negritas
    paragraph.runs[0].bold = True  # Acceder al primer "run" y aplicar negritas
    set_vertical_alignment(cell_titulo, "center")




    cell_separador = table.cell(4, 0)
    cell_separador.merge(table.cell(4, 1))
    remove_cell_borders(table.cell(4, 0), ['left', 'bottom', 'right', 'top'])

    
    # Aplicar sombreado personalizado a la celda fusionada
    shading_color = RGBColor(144, 10, 80)
    shading_color2 = RGBColor(168, 208, 141)

    # Convertir RGB a hexadecimal manualmente
    def rgb_to_hex(rgb_color):
        return '{:02X}{:02X}{:02X}'.format(rgb_color[0], rgb_color[1], rgb_color[2])

    hex_color = rgb_to_hex(shading_color)  # Obtener el valor hexadecimal
    hex_color2 = rgb_to_hex(shading_color2)  # Obtener el valor hexadecimal

    # Aplicar el sombreado
    shading_element = OxmlElement('w:shd')
    shading_element.set(qn('w:fill'), hex_color)  # Usar el valor HEX sin el prefijo "#"
    cell_titulo._element.tcPr.append(shading_element)
    
    remove_cell_borders(table.cell(3, 0), ['right', 'left', 'top'])

    # --- Información del paciente ---
    table_paciente = doc.add_table(rows=4, cols=2)
    table_paciente.style = 'Table Grid'

    # Ajustar el alto de las filas
    for row in table_paciente.rows:
        row.height = Inches(0.3)  # Establece el alto de las filas a 0.5 pulgadas

    # Ajustar el ancho de las columnas
    for col in table_paciente.columns:
        for cell in col.cells:
            cell.width = Inches(3)  # Establece el ancho de las columnas a 3 pulgadas

    # Llenar la tabla con etiquetas y valores
    table_paciente.cell(0, 0).text = 'Nombre'
    table_paciente.cell(0, 1).text = datos["nombre"]
    table_paciente.cell(1, 0).text = 'Edad'
    table_paciente.cell(1, 1).text = datos["edad"]
    table_paciente.cell(2, 0).text = 'Fecha de Nacimiento'
    table_paciente.cell(2, 1).text = datos["fecha_nacimiento"]
    table_paciente.cell(3, 0).text = 'Diagnóstico'
    table_paciente.cell(3, 1).text = datos["diagnostico"]

    # Aplicar negritas a las etiquetas
    for i in range(4):  # Recorre las filas de la tabla
        cell = table_paciente.cell(i, 0)  # Accede a la primera columna (etiquetas)
        paragraph = cell.paragraphs[0]  # Obtiene el primer párrafo de la celda
        paragraph.runs[0].bold = True  # Aplica negritas al primer run

    # Alinear el contenido de las celdas a la izquierda en horizontal y al centro en vertical
    for row in table_paciente.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_vertical_alignment(cell, "center")
#    Agregar un espacio intermedio entre las tablas
    doc.add_paragraph()  # Agrega un párrafo vacío
    # --- Área de intervención ---
    table_intervencion = doc.add_table(rows=4, cols=2)
    table_intervencion.style = 'Table Grid'

    # Ajustar el alto de las filas
    for row in table_intervencion.rows:
        row.height = Inches(0.5)  # Establece el alto de las filas a 0.5 pulgadas

    # Ajustar el ancho de las columnas
    for col in table_intervencion.columns:
        for cell in col.cells:
            cell.width = Inches(3)  # Establece el ancho de las columnas a 3 pulgadas

    # Llenar la tabla
    table_intervencion.cell(0, 0).text = 'Área de intervención'
    table_intervencion.cell(0, 1).text = datos["area_intervencion"]
    table_intervencion.cell(1, 0).text = 'Periodo de intervencion'
    table_intervencion.cell(1, 1).text = datos["periodo_intervencion"]
    table_intervencion.cell(2, 0).text = 'Número de terapias recibidas'
    table_intervencion.cell(2, 1).text = datos["terapias_recibidas"]
    table_intervencion.cell(3, 0).text = 'Número de faltas'
    table_intervencion.cell(3, 1).text = datos["faltas"]

    # Aplicar negritas a las etiquetas
    for i in range(4):  # Recorre las filas de la tabla
        cell = table_intervencion.cell(i, 0)  # Accede a la primera columna (etiquetas)
        paragraph = cell.paragraphs[0]  # Obtiene el primer párrafo de la celda
        paragraph.runs[0].bold = True  # Aplica negritas al primer run

    # Alinear el contenido de las celdas a la izquierda en horizontal y al centro en vertical
    for row in table_intervencion.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_vertical_alignment(cell, "center")


    # --- Objetivos iniciales ---
    doc.add_paragraph()  # Espacio vacío
    p = doc.add_paragraph()
    p.add_run('Objetivos iniciales').bold = True
    

    # Añadir objetivos iniciales como lista
    for objetivo in datos["objetivos_iniciales"].split('\n'):
        doc.add_paragraph(objetivo, style='List Bullet')

    # --- Avance de los objetivos planteados ---
    doc.add_paragraph()  # Espacio vacío
    p = doc.add_paragraph()
    p.add_run('Avance de los objetivos planteados').bold = True

    # Dividir el texto en líneas
    lineas = datos["avance_objetivos"].split('\n')

    # Agregar la primera línea como un párrafo normal (sin viñeta)
    if lineas:  # Verificar que haya al menos una línea
        doc.add_paragraph(lineas[0].strip())  # Primera línea sin viñeta

    # Agregar el resto de las líneas como párrafos con viñeta
    for linea in lineas[1:]:
        if linea.strip():  # Ignorar líneas vacías
            doc.add_paragraph(linea, style='List Bullet')

    # --- Nuevos objetivos ---
    doc.add_paragraph()  # Espacio vacío
    p = doc.add_paragraph()
    p.add_run('Nuevos objetivos').bold = True

    # Añadir nuevos objetivos como lista
    for objetivo in datos["nuevos_objetivos"].split('\n'):
        doc.add_paragraph(objetivo, style='List Bullet')

    # --- Seguimiento ---
    doc.add_paragraph()  # Espacio vacío
    p = doc.add_paragraph()
    p.add_run('Seguimiento').bold = True

    # Subsección: Observaciones
    p = doc.add_paragraph()
    p.add_run('Observaciones:').bold = True
    doc.add_paragraph(datos["observaciones"])

    # Subsección: Tratamiento
    p = doc.add_paragraph()
    p.add_run('Tratamiento:').bold = True
    doc.add_paragraph(datos["tratamiento"])

# --- Sugerencias para casa ---
    section_heading = doc.add_paragraph('Sugerencias Para Casa')
    section_heading.style = doc.styles['Heading 1']
    section_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro
    sugerencias_casa = datos.get("sugerencias_casa", {})

    def procesar_contenido(contenido):
        """Procesa diferentes tipos de contenido recursivamente"""
        if isinstance(contenido, dict):
            # Convertir diccionarios a formato clave-valor
            items = []
            for k, v in contenido.items():
                items.append(f"{k}: {procesar_contenido(v)}")
            return '\n'.join(items)
        elif isinstance(contenido, list):
            # Procesar cada elemento de la lista
            return '\n'.join(f'• {procesar_contenido(item)}' for item in contenido)
        return str(contenido)

    # 1. Procesar categorías principales
    if "Categorías" in sugerencias_casa:
        for categoria in sugerencias_casa["Categorías"]:
            # Encabezado de categoría
            category_heading = doc.add_paragraph(style='Normal')
            category_run = category_heading.add_run(f"{categoria.get('Icono', '')} {categoria.get('Nombre', '')}")
            category_run.bold = True
            category_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Procesar sugerencias individuales
            for sugerencia in categoria.get("Sugerencias", []):
                # Objetivos
                if objetivos := sugerencia.get("Objetivos"):
                    obj_heading = doc.add_paragraph(style='Normal')
                    obj_run = obj_heading.add_run('Objetivos: ')
                    obj_run.bold = True
                    obj_run.font.color.rgb = RGBColor(0, 0, 0)
                    for obj in objetivos:
                        p = doc.add_paragraph(procesar_contenido(obj), style='List Bullet 2')
                        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                
                # Campos dinámicos
                campos_especiales = [
                    ('Fundamento ABA', 'Explicación'),
                    ('Escalabilidad', 'Registro'),
                    ('Progresión', 'Técnica')
                ]
                
                for campo1, campo2 in campos_especiales:
                    valor1 = sugerencia.get(campo1)
                    valor2 = sugerencia.get(campo2)
                    
                    if valor1 or valor2:
                        field_paragraph = doc.add_paragraph(style='Normal')
                        if valor1:
                            field_run = field_paragraph.add_run(f"{campo1}: ")
                            field_run.bold = True
                            field_run.font.color.rgb = RGBColor(0, 0, 0)
                            field_paragraph.add_run(procesar_contenido(valor1))
                        
                        if valor2:
                            field_run = field_paragraph.add_run("\n" + f"{campo2}: ")
                            field_run.bold = True
                            field_run.font.color.rgb = RGBColor(0, 0, 0)
                            field_paragraph.add_run(procesar_contenido(valor2))
                
                # Ejemplos y materiales
                for seccion in ['Ejemplos', 'Materiales']:
                    if items := sugerencia.get(seccion):
                        sec_paragraph = doc.add_paragraph(style='Normal')
                        sec_run = sec_paragraph.add_run(f"{seccion}: ")
                        sec_run.bold = True
                        sec_run.font.color.rgb = RGBColor(0, 0, 0)
                        contenido = procesar_contenido(items)
                        for line in contenido.split('\n'):
                            p = doc.add_paragraph(line, style='List Bullet 2')
                            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                

    # 2. Recomendaciones generales
    if recomendaciones := sugerencias_casa.get("Recomendaciones Generales"):
        rec_heading = doc.add_paragraph('Recomendaciones Generales', style='Heading 1')
        rec_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        for titulo, contenido in recomendaciones.items():
            p = doc.add_paragraph(style='List Bullet 2')
            p.add_run(f"{titulo}: ").bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            p.add_run(procesar_contenido(contenido)).font.color.rgb = RGBColor(0, 0, 0)
            
    doc.add_paragraph()  # Espacio vacío

    # Crear un párrafo para "Elaborado por" y centrarlo
    p_elaborado = doc.add_paragraph(f'Elaborado por: {datos["elaborado_por"]}')
    p_elaborado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Crear un párrafo para "Cédula" y centrarlo
    p_cedula = doc.add_paragraph(f'Cédula: {datos["cedula"]}')
    p_cedula.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar el documento
    doc.save(nombre_archivo)
