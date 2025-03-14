import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Cargar los datos del archivo JSON
with open('expedientes.json', 'r', encoding='utf-8') as file:
    pacientes = json.load(file)

# Seleccionar un paciente (por ejemplo, el primero)
paciente = pacientes[0]

# Crear un nuevo documento de Word
doc = Document()

# Establecer el estilo de fuente predeterminado
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Crear una tabla con 3 filas y 4 columnas
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'

# Llenar la tabla con el logo, fecha, unidad y expediente
table.cell(0, 0).text = 'LOGO'
table.cell(0, 1).text = 'LOGO'
table.cell(0, 2).text = 'Fecha: ' + paciente.get('fecha', '')
table.cell(0, 3).text = 'Unidad: ' + paciente.get('unidad', '')
table.cell(1, 2).text = 'Expediente: ' + paciente.get('expediente', '')

# Fusionar celdas para el mensaje "INFORME DE SEGUIMIENTO"
cell = table.cell(2, 0)
cell.merge(table.cell(2, 3))
cell.text = 'INFORME DE SEGUIMIENTO'
cell.paragraphs[0].alignment = 1  # Centrar el texto

# Añadir una fila vacía
doc.add_paragraph()

# Añadir la información del paciente
doc.add_paragraph('Nombre: ' + paciente.get('nombre', ''))
doc.add_paragraph('Edad: ' + str(paciente.get('edad', '')))
doc.add_paragraph('Fecha de nacimiento: ' + paciente.get('fecha_nacimiento', ''))
doc.add_paragraph('Diagnóstico: ' + paciente.get('diagnostico', ''))

# Añadir una fila vacía
doc.add_paragraph()

# Añadir información de intervención
doc.add_paragraph('Área de intervención: ' + paciente.get('area_intervencion', ''))
doc.add_paragraph('Periodo de intervención: ' + paciente.get('periodo_intervencion', ''))
doc.add_paragraph('Número de terapias recibidas: ' + str(paciente.get('terapias_recibidas', '')))
doc.add_paragraph('Número de faltas: ' + str(paciente.get('faltas', '')))

# Añadir una fila vacía
doc.add_paragraph()

# Añadir objetivos iniciales
p = doc.add_paragraph()
p.add_run('Objetivos iniciales').bold = True
p.add_run().font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Verde aguacate

for objetivo in paciente.get('objetivos_iniciales', []):
    doc.add_paragraph(objetivo, style='List Bullet')

# Añadir avance de los objetivos planteados mediante IA
doc.add_paragraph('Avance de los objetivos planteados mediante IA: ' + paciente.get('avance_objetivos', ''))

# Añadir nuevos objetivos mediante IA
p = doc.add_paragraph()
p.add_run('Nuevos objetivos').bold = True
p.add_run().font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Verde aguacate

for objetivo in paciente.get('nuevos_objetivos', []):
    doc.add_paragraph(objetivo, style='List Bullet')

# Añadir seguimiento, recomendaciones generales, observaciones, tratamiento y sugerencias para casa
doc.add_paragraph('Seguimiento: ' + paciente.get('seguimiento', ''))
doc.add_paragraph('Recomendaciones generales: ' + paciente.get('recomendaciones', ''))
doc.add_paragraph('Observaciones: ' + paciente.get('observaciones', ''))
doc.add_paragraph('Tratamiento: ' + paciente.get('tratamiento', ''))
doc.add_paragraph('Sugerencias para casa: ' + paciente.get('sugerencias_casa', ''))

# Añadir elaborado por y cédula desde configuración
doc.add_paragraph('Elaborado por: ' + paciente.get('elaborado_por', ''))
doc.add_paragraph('Cédula: ' + paciente.get('cedula', ''))

# Guardar el documento
doc.save('Informe_Seguimiento_Ejemplo.docx')
