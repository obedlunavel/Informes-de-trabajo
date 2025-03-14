from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import tkinter as tk
from tkinter import ttk

def generar_documento(datos):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Encabezado
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    cell_logo = table.cell(0, 0)
    cell_logo.merge(table.cell(1, 1))
    cell_logo.text = 'LOGO'

    cell_info = table.cell(0, 2)
    cell_info.merge(table.cell(1, 3))
    p = cell_info.add_paragraph(f'FECHA: {datos["fecha"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = cell_info.add_paragraph(f'UNIDAD: {datos["unidad"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = cell_info.add_paragraph(f'NO. EXPEDIENTE: {datos["expediente"]}')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cell_titulo = table.cell(2, 0)
    cell_titulo.merge(table.cell(2, 3))
    cell_titulo.text = 'INFORME DE SEGUIMIENTO'
    cell_titulo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Información del paciente
    table_paciente = doc.add_table(rows=4, cols=2)
    table_paciente.style = 'Table Grid'
    table_paciente.cell(0, 0).text = 'NOMBRE'
    table_paciente.cell(0, 1).text = datos["nombre"]
    table_paciente.cell(1, 0).text = 'EDAD'
    table_paciente.cell(1, 1).text = datos["edad"]
    table_paciente.cell(2, 0).text = 'FECHA DE NACIMIENTO'
    table_paciente.cell(2, 1).text = datos["fecha_nacimiento"]
    table_paciente.cell(3, 0).text = 'DIAGNÓSTICO'
    table_paciente.cell(3, 1).text = datos["diagnostico"]

    # Área de intervención
    table_intervencion = doc.add_table(rows=4, cols=2)
    table_intervencion.style = 'Table Grid'
    table_intervencion.cell(0, 0).text = 'ÁREA DE INTERVENCIÓN'
    table_intervencion.cell(0, 1).text = datos["area_intervencion"]
    table_intervencion.cell(1, 0).text = 'PERIODO DE INTERVENCIÓN'
    table_intervencion.cell(1, 1).text = datos["periodo_intervencion"]
    table_intervencion.cell(2, 0).text = 'NÚMERO DE TERAPIAS RECIBIDAS'
    table_intervencion.cell(2, 1).text = datos["terapias_recibidas"]
    table_intervencion.cell(3, 0).text = 'NÚMERO DE FALTAS'
    table_intervencion.cell(3, 1).text = datos["faltas"]

    # Objetivos iniciales
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('OBJETIVOS INICIALES').bold = True
    for objetivo in datos["objetivos_iniciales"].split('\n'):
        doc.add_paragraph(objetivo, style='List Bullet')

    # Avance de los objetivos planteados
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('AVANCE DE LOS OBJETIVOS PLANTEADOS').bold = True
    doc.add_paragraph(datos["avance_objetivos"])

    # Nuevos objetivos
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('NUEVOS OBJETIVOS').bold = True
    for objetivo in datos["nuevos_objetivos"].split('\n'):
        doc.add_paragraph(objetivo, style='List Bullet')

    # Seguimiento
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('SEGUIMIENTO').bold = True
    doc.add_paragraph(datos["seguimiento"])

    # Elaborado por
    doc.add_paragraph()
    doc.add_paragraph(f'Elaborado por: {datos["elaborado_por"]}')
    doc.add_paragraph(f'Cédula: {datos["cedula"]}')

    doc.save(f'IS_{datos["nombre"].replace(" ", "_")}_LIBERADO_Replica.docx')

def on_submit():
    datos = {
        "fecha": fecha_entry.get(),
        "unidad": unidad_entry.get(),
        "expediente": expediente_entry.get(),
        "nombre": nombre_entry.get(),
        "edad": edad_entry.get(),
        "fecha_nacimiento": fecha_nacimiento_entry.get(),
        "diagnostico": diagnostico_entry.get(),
        "area_intervencion": area_intervencion_entry.get(),
        "periodo_intervencion": periodo_intervencion_entry.get(),
        "terapias_recibidas": terapias_recibidas_entry.get(),
        "faltas": faltas_entry.get(),
        "objetivos_iniciales": objetivos_iniciales_entry.get("1.0", tk.END),
        "avance_objetivos": avance_objetivos_entry.get("1.0", tk.END),
        "nuevos_objetivos": nuevos_objetivos_entry.get("1.0", tk.END),
        "seguimiento": seguimiento_entry.get("1.0", tk.END),
        "elaborado_por": elaborado_por_entry.get(),
        "cedula": cedula_entry.get()
    }
    generar_documento(datos)
    root.destroy()

root = tk.Tk()
root.title("Generador de Informe de Seguimiento")

# Crear campos de entrada
tk.Label(root, text="Fecha:").grid(row=0, column=0)
fecha_entry = tk.Entry(root)
fecha_entry.grid(row=0, column=1)

tk.Label(root, text="Unidad:").grid(row=1, column=0)
unidad_entry = tk.Entry(root)
unidad_entry.grid(row=1, column=1)

tk.Label(root, text="Número de Expediente:").grid(row=2, column=0)
expediente_entry = tk.Entry(root)
expediente_entry.grid(row=2, column=1)

tk.Label(root, text="Nombre del Paciente:").grid(row=3, column=0)
nombre_entry = tk.Entry(root)
nombre_entry.grid(row=3, column=1)

tk.Label(root, text="Edad:").grid(row=4, column=0)
edad_entry = tk.Entry(root)
edad_entry.grid(row=4, column=1)

tk.Label(root, text="Fecha de Nacimiento:").grid(row=5, column=0)
fecha_nacimiento_entry = tk.Entry(root)
fecha_nacimiento_entry.grid(row=5, column=1)

tk.Label(root, text="Diagnóstico:").grid(row=6, column=0)
diagnostico_entry = tk.Entry(root)
diagnostico_entry.grid(row=6, column=1)

tk.Label(root, text="Área de Intervención:").grid(row=7, column=0)
area_intervencion_entry = tk.Entry(root)
area_intervencion_entry.grid(row=7, column=1)

tk.Label(root, text="Período de Intervención:").grid(row=8, column=0)
periodo_intervencion_entry = tk.Entry(root)
periodo_intervencion_entry.grid(row=8, column=1)

tk.Label(root, text="Número de Terapias Recibidas:").grid(row=9, column=0)
terapias_recibidas_entry = tk.Entry(root)
terapias_recibidas_entry.grid(row=9, column=1)

tk.Label(root, text="Número de Faltas:").grid(row=10, column=0)
faltas_entry = tk.Entry(root)
faltas_entry.grid(row=10, column=1)

tk.Label(root, text="Objetivos Iniciales:").grid(row=11, column=0)
objetivos_iniciales_entry = tk.Text(root, height=5, width=30)
objetivos_iniciales_entry.grid(row=11, column=1)

tk.Label(root, text="Avance de los Objetivos Planteados:").grid(row=12, column=0)
avance_objetivos_entry = tk.Text(root, height=5, width=30)
avance_objetivos_entry.grid(row=12, column=1)

tk.Label(root, text="Nuevos Objetivos:").grid(row=13, column=0)
nuevos_objetivos_entry = tk.Text(root, height=5, width=30)
nuevos_objetivos_entry.grid(row=13, column=1)

tk.Label(root, text="Seguimiento:").grid(row=14, column=0)
seguimiento_entry = tk.Text(root, height=5, width=30)
seguimiento_entry.grid(row=14, column=1)

tk.Label(root, text="Elaborado por:").grid(row=15, column=0)
elaborado_por_entry = tk.Entry(root)
elaborado_por_entry.grid(row=15, column=1)

tk.Label(root, text="Cédula:").grid(row=16, column=0)
cedula_entry = tk.Entry(root)
cedula_entry.grid(row=16, column=1)

# Botón de envío
submit_button = tk.Button(root, text="Generar Documento", command=on_submit)
submit_button.grid(row=17, column=0, columnspan=2)

root.mainloop()
