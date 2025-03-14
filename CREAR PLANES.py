import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.shared import Pt, RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import tkinter.simpledialog as simpledialog
import sys


# © 2025, Obed Luna Velázquez. Todos los derechos reservados.
# Este programa y su código fuente están protegidos por las leyes de derechos de autor.
# Prohibida su distribución y/o modificación sin autorización expresa del autor.
def resource_path(relative_path):
    """Obtiene la ruta correcta de un archivo, incluso dentro del ejecutable."""
    try:
        # PyInstaller usa este directorio temporal
        base_path = sys._MEIPASS
    except AttributeError:
        # Ruta normal si no está empaquetado
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Diccionario con los programas de intervención
programas = {
    1: {
        "Nombre": "Respuesta al nombre",
        "Objetivo": "Que el alumno(a) dirija y enfoque su atención a la cara de la persona al escuchar su nombre estando fuera de su campo visual.",
        "Procedimiento": "Enfoque 1: Sentar al alumno(a) frente al cuidador y tomar un estímulo de interés para el alumno(a), dirigirlo hacia el entrecejo del cuidador mientras se le llama por su nombre. Enfoque 2: Participar en una actividad de interés del alumno(a), llamarlo por su nombre durante la interacción.",
        "Ayudas": "Ayuda visual: Mover dentro del campo visual el estímulo de interés y guiarlo hacia el entrecejo del terapeuta mientras se dice su nombre. Ayuda física: Tomar ambas manos del alumno(a) y ponerlas en cara del terapeuta mientras se dice su nombre."
    },
    2: {
        "Nombre": "Mírame",
        "Objetivo": "Que el alumno(a) dirija y enfoque su atención a la cara de la persona que lo llama realizando contacto visual.",
        "Procedimiento": "Tomar un estímulo de interés y colocarlo en el entrecejo mientras se dice “mírame” o “nombre + mírame”.",
        "Ayudas": "Ayuda visual: Mover dentro del campo visual el estímulo de interés y guiarlo hacia el entrecejo del terapeuta. Ayuda física: Tomar ambas manos del niño y ponerlas en cara del terapeuta para que el niño dirija su mirada."
    },
    3: {
        "Nombre": "Mira",
        "Objetivo": "Que el alumno(a) dirija su atención hacia el estímulo que se le señala.",
        "Procedimiento": "Enfoque 1: Colocar un objeto de interés cerca del alumno(a) y decir “Mira el perro” mientras se señala con el dedo índice. Enfoque 2: Participar en una actividad de interés del alumno(a) y decir “mira la oveja” mientras se señala.",
        "Ayudas": "Ayuda visual: Mover un estímulo de su interés dentro de su campo visual y redirigir su atención hacía el oso de peluche."
    },
    4: {
        "Nombre": "Imitaciones motoras gruesas",
        "Subobjetivo": [
            "Golpear mesa cercana",
            "levantar brazos",
            "tocarse cabeza",
            "tocarse estómago",
            "tocarse rodillas",
            "tocarse pies",
            "patalear",
            "lanzar besos",
            "decir adiós (hola) con la mano",
            "decir sí con cabeza",
            "decir no con cabeza",
            "aplaudir"
        ],
        "Objetivo": "Que mueva manos y pies replicando el modelo visual.",
        "Procedimiento": "Enfoque 1: Sentar al alumno(a) frente al cuidador, llamarlo por su nombre y decir “haz esto” mientras se realizan movimientos con las manos o pies. Enfoque 2: Cantar canciones de interés del alumno(a) integrando los objetivos como parte de la coreografía.",
        "Ayudas": "Ayuda física: Tomar ambas manos del alumno(a) y levantarlas. Ayuda visual: Tomar un objeto de su interés y colocarlo por encima de la cabeza del alumno(a) para que intente tomarlo."
    },
    5: {
        "Nombre": "Imitaciones motoras finas",
        "Objetivo": "Que mueva manos y dedos replicando el modelo visual.",
        "Procedimiento": "Sentar al alumno(a) de frente, llamarlo por su nombre y decir “haz esto” mientras se realizan movimientos finos como frotar las manos.",
        "Ayudas": "Ayuda física: Tomar ambas manos del alumno(a) y hacer los movimientos de frotar las palmas (hacia enfrente y hacia atrás) junto con él."
    },
    6: {
        "Nombre": "Imitación de movimientos orofaciales",
        "Objetivo": "Que el alumno(a) mueva los músculos de la boca replicando el modelo visual. Ej. sacar la lengua.",
        "Procedimiento": "1. Llamar al alumno(a) por su nombre: 'nombre + mírame' para que realice contacto visual y anticipar la siguiente actividad. 2. Decir 'haz esto' mientras se saca la lengua. 3. Si lo hace correctamente se refuerza. 4. Si no lo hace o no hay aproximaciones, se le brindan las ayudas necesarias para cumplir con el objetivo del programa.",
        "Ayudas": "Ayuda física: Con ayuda de un abatelenguas, colocar el abatelenguas en la lengua del alumno(a) realizando movimientos para que saque la lengua. Ayuda visual: Colocar frente a su boca una paleta motivando a que saque la lengua para lamerla."
    },
    7: {
        "Nombre": "Seguimiento de instrucciones básicas",
        "Objetivo": "Que el alumno(a) ejecute la instrucción que se le verbaliza. Ej. 'Siéntate'.",
        "Procedimiento": "1. Sentarse frente al alumno(a). 2. Colocar una silla detrás de él/ella. 3. Dar la instrucción 'Siéntate'. 4. Si realiza la acción, se refuerza. 5. Si no hay respuesta o es incorrecta, se le brindan las ayudas necesarias para cumplir con el objetivo del programa.",
        "Ayudas": "Ayuda física: Tomar al alumno(a) de sus brazos y ayudarle a sentarse al mismo tiempo que se le da la instrucción. Otra ayuda física: Tocar sus hombros al mismo tiempo en el que se le da la instrucción."
    },
    8: {
        "Nombre": "Señalar con índice",
        "Objetivo": "Que el alumno(a) señale objetos de interés con el dedo índice.",
        "Procedimiento": "1. Colocar objetos o imágenes de interés del alumno(a) dentro de su campo visual. 2. Apoyar físicamente para que el alumno(a) coloque su mano en posición de señalamiento. 3. Señalar junto con el alumno(a) los objetos al tiempo que se le nombra el material reforzando el vocabulario.",
        "Ayudas": "Ayuda física: Tomar su mano y ayudarlo a poner el dedo índice en posición de señalamiento hacia el objeto deseado. Si ya domina la habilidad, mover su mano hacia el objeto que va a señalar."
    },
    9: {
        "Nombre": "Responder sí y no con la cabeza",
        "Objetivo": "Que el alumno(a) responda correctamente por medio de un gesto (sí/no) a preguntas de 'sí' o 'no'.",
        "Procedimiento": "1. Presentarle algo que le guste o sea de su agrado. 2. Realizar la pregunta '¿Quieres?' 3. Si responde correctamente, refuérzalo. 4. Si no responde o es incorrecto, repetir la actividad y brindarle las ayudas necesarias.",
        "Ayudas": "Ayuda visual: Modelar el gesto de 'sí' moviendo la cabeza de arriba hacia abajo. Ayuda física: Tomar la cabeza del alumno(a) y moverla lentamente hacia arriba y hacia abajo."
    },
    10: {
        "Nombre": "Reconocimiento de esquema corporal",
        "Objetivo": "Que el alumno(a) reconozca las partes del cuerpo.",
        "Procedimiento": "1. Sentar al alumno(a) de frente. 2. Hacer la pregunta '¿Dónde está tu cabeza?' o 'toca tu cabeza'. 3. Si lo hace correctamente, se refuerza. 4. Si no lo hace, se le brindan las ayudas necesarias para tener éxito en el objetivo del programa.",
        "Ayudas": "Ayuda física: Tomar la mano del alumno(a) dirigiéndola hacia su cabeza. Otra ayuda física: Tocar suavemente los codos del alumno(a) dirigiéndolo hacia la cabeza."
    },
    11: {
        "Nombre": "Emparejamiento de estímulos visuales",
        "Objetivo": "Que el alumno(a) coloque dos estímulos iguales juntos siguiendo la instrucción 'pon con el mismo'.",
        "Procedimiento": "Colocar solamente un objeto o imagen sobre la mesa. Proporcionar el mismo objeto al alumno(a) y decir 'pon con igual'. Apoyar al alumno(a) a colocar el objeto junto al que ya está en la mesa. Si lo hace correctamente, reforzar. Si no hay respuesta o el alumno(a) intenta tomar el objeto, brindar las ayudas necesarias. Aumentar la cantidad de estímulos conforme se va dominando el objetivo.",
        "Ayudas": "Ayuda física: Tomar la mano del alumno(a), colocarla sobre uno de los objetos y guiarla para llevar el objeto al lugar correcto. Ayuda visual: Señalar el contenedor donde debe colocar el estímulo mientras se dice 'pon con igual'."
    },
    12: {
        "Nombre": "Vocabulario receptivo",
        "Objetivo": "Que el alumno(a) identifique señalando o entregando la imagen del objeto que se le pide.",
        "Procedimiento": "1. Nombrar la categoría con la que se va a trabajar (ej. familia). 2. Mostrar las tarjetas nombrando cada imagen. 3. Colocar una sola imagen frente al alumno(a). 4. Preguntar '¿Dónde está la abuela?' (por ejemplo). 5. Si el alumno(a) señala correctamente, reforzar. Si no lo hace, brindar las ayudas necesarias. Aumentar el número de estímulos conforme el alumno(a) domine el objetivo.",
        "Ayudas": "Ayuda física: Tomar la mano del alumno(a) y colocarla sobre el estímulo correcto. Ayuda visual: Señalar con el dedo índice el estímulo correcto. Ayuda de terceros: Un acompañante se coloca detrás del alumno(a) y lo ayuda a tomar el estímulo correcto."
    },
    13: {
        "Nombre": "Seguimiento de instrucciones con objetos",
        "Objetivo": "Que el alumno(a) siga una instrucción de un solo paso.",
        "Procedimiento": "Sentar al alumno(a) de frente y captar su atención. Realizar una actividad de interés del alumno(a) (ej. tocar el tambor). Dar la instrucción 'para' y hacer que la actividad se detenga por breves segundos. Si lo hace correctamente, reforzar. Si no, repetir la actividad brindando las ayudas necesarias.",
        "Ayudas": "Ayuda física: Tomar las manos del alumno(a) y ayudarle a detener la acción. Ayuda de terceros: Un acompañante puede ayudar a detener la acción tomando las manos del alumno(a)."
    },
    14: {
        "Nombre": "Seguimiento de instrucciones complejas: encadenadas de 2-3 pasos",
        "Objetivo": "Que el alumno(a) siga la instrucción realizando ambas acciones en el orden indicado. Ej. 'Aplaude y sube tus brazos'.",
        "Procedimiento": "1. Sentar al alumno(a) de frente y captar su atención. 2. Dar la instrucción 'aplaude y sube tus brazos'. 3. Esperar que el alumno(a) realice la actividad. 4. Si lo hace correctamente, reforzar. Si no, brindar las ayudas necesarias.",
        "Ayudas": "Ayuda física: Tomar las manos del alumno(a) y apoyarlo a realizar las acciones en el orden indicado. Ayuda visual: Modelar las acciones después de dar la instrucción."
    },    
    15: {
        "Nombre": "Seguimiento de instrucciones complejas por trayectoria",
        "Objetivo": "Que el alumno(a) sea capaz de realizar una instrucción que implique caminar cierta distancia. Ej. tirar la basura, apagar/prender la luz.",
        "Procedimiento": "Sentar al alumno(a) de frente y captar su atención. Dar la instrucción (ej. 'tíralo' o 'a la basura'). Si el alumno(a) realiza la instrucción correctamente, reforzar. Si no, brindar las ayudas necesarias. Se comienza con el bote de basura cerca y se va alejando conforme el alumno(a) domina la instrucción.",
        "Ayudas": "Ayuda física: Tomar la mano del alumno(a) y dirigirlo al bote de basura ayudándolo a tirar la basura. Ayuda visual: Dirigir su mirada hacia el bote de basura señalándolo con el dedo índice."
    },
    16: {
        "Nombre": "Seguimiento de instrucciones complejas por trayectoria de 2-3 órdenes",
        "Objetivo": "Que el alumno(a) realice las instrucciones que se le darán en el orden indicado.",
        "Procedimiento": "1. Que el alumno(a) haga contacto ocular antes de darle las instrucciones. 2. Darle las instrucciones al alumno(a) ‘cierra la puerta y apaga la luz’. 3. Se espera que el alumno(a) se levante de su silla y realice las acciones en el orden indicado. 4. Si lo realiza correctamente se reforzará. 5. Si no lo realiza o lo hace incorrecto, se brindarán las ayudas necesarias para cumplir con el objetivo.",
        "Ayudas": "Ayuda física: Se toma al alumno(a) de la mano y lo ayudamos a realizar las acciones en el orden indicado. Ayuda visual: Dirigir su mirada utilizando el dedo índice hacia la puerta y el interruptor de la luz."
    },
    17: {
        "Nombre": "Imitación de onomatopeyas",
        "Objetivo": "Imitar sonidos de objetos, cuerpo, entre otros.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y mantener su atención. Mostrarle la imagen de un animal (ej. vaca) y hacerle el sonido ‘muuuuu’. Repetir el sonido varias veces y darle tiempo al alumno(a) para que lo repita. Si responde correctamente, reforzar. Si no hay respuesta o aproximación, brindar las ayudas necesarias para tener éxito en el objetivo.",
        "Ayudas": "Ayuda verbal: Decir una sílaba y esperar a que el alumno(a) complete el sonido. Ayudas visuales: Mostrar imágenes relacionadas al sonido que se realiza."
    },
    18: {
        "Nombre": "Ejercicios orofaciales",
        "Objetivo": "Estimular la conciencia oral y la coordinación de los órganos fonoarticuladores.",
        "Procedimiento": "Colocarse frente al alumno(a) y captar su atención previamente. Dar la instrucción ‘haz así’ seguido de un movimiento orofacial, por ejemplo, sacar la lengua. Esperar que el niño(a) imite la acción. Si responde correctamente, reforzar. Si la respuesta es incorrecta, repetir la instrucción.",
        "Ayudas": "Ayuda visual: Apoyarse de un espejo. Ayuda física: Colocar las manos en los cachetes para provocar el movimiento adecuado."
    },
    19: {
        "Nombre": "Vocales",
        "Objetivo": "Mencionar e imitar vocales.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y mantener su atención. Mostrar la vocal ‘A’ y vocalizar. Repetir el sonido varias veces y darle tiempo al alumno(a) para que lo repita. Si responde correctamente, reforzar. Si no hay respuesta o aproximación, brindar las ayudas necesarias.",
        "Ayudas": "Ayudas visuales: Tarjetas o imágenes de las vocales. Ayuda física sensorial: Colocar espuma en la mano y realizar el movimiento de la forma de la vocal. Ayuda física: Abate lenguas para realizar el correcto acomodo de la boca o lengua."
    },
    20: {
        "Nombre": "Sílabas",
        "Objetivo": "Imitar y mencionar cadenas silábicas comenzando por las más comunes e ir aumentando la complejidad de los fonemas.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y mantener su atención. Vocalizar ‘mamama’. Repetir el sonido varias veces y darle tiempo al alumno(a) para que lo repita. Si responde correctamente, reforzar. Si no hay respuesta ni aproximaciones, brindar las ayudas necesarias.",
        "Ayudas": "Ayuda visual: Imágenes de tarjetas con las sílabas. Ayuda física: Abate lenguas para realizar el correcto acomodo de la boca o lengua."
    },
    21: {
        "Nombre": "Ecóicas",
        "Objetivo": "Repetir el vocabulario que el terapeuta le menciona.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y mantener su atención. Se le muestra la tarjeta de un animal y se nombra. Repetir el nombre varias veces y dar tiempo al alumno(a) para que lo repita. Si responde correctamente, reforzar. Si no hay respuesta o aproximación, se brindan las ayudas necesarias.",
        "Ayudas": "Ayuda verbal: nombrar la primera sílaba de la palabra y esperar a que el alumno(a) la complete. Ayuda física: tomar la mano del alumno(a) y colocarla en el objeto para guiar su atención."
    },
    22: {
        "Nombre": "Mandos (palabras para pedir)",
        "Objetivo": "Que el alumno(a) solicite verbalmente el objeto que desea.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y captar su atención. Mostrar al alumno(a) un objeto o juguete de interés y mantenerlo en su campo visual. Incentivar al alumno(a) a que lo nombre o lo pida verbalmente.",
        "Ayudas": "Ayuda verbal: decir la primera sílaba o palabra completa, y si el alumno(a) lo repite, se le entrega el objeto. Ayuda física: guiar su mano al objeto mientras se nombra."
    },
    23: {
        "Nombre": "Tactos",
        "Objetivo": "El alumno(a) deberá nombrar la imagen u objeto que se le presente.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta, presentar cuatro imágenes una por una mientras se nombran, y pedir al alumno(a) que las nombre señalando.",
        "Ayudas": "Ayuda verbal: decir la primera sílaba de la palabra. Ayuda visual: mostrar una tarjeta a la vez. Ayuda física: guiar la mano del alumno(a) hacia la tarjeta."
    },
    24: {
        "Nombre": "Vocabulario",
        "Objetivo": "Presentar y reconocer el vocabulario a través de campos semánticos.",
        "Procedimiento": "Presentar el vocabulario por categorías mediante tarjetas o material didáctico. Pedir al alumno(a) que señale o nombre el material, reforzar si lo hace correctamente.",
        "Ayudas": "Ayuda verbal: mencionar la primera parte de la palabra como recordatorio. Ayuda de terceros: otro adulto puede susurrar la respuesta correcta al oído del alumno(a)."
    },
    25: {
        "Nombre": "Intraverbales",
        "Objetivo": "Responder a preguntas de información personal.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y captar su atención. Comenzar con preguntas sencillas como ‘¿Cómo te llamas?’ y avanzar progresivamente.",
        "Ayudas": "Ayuda verbal: pronunciar la primera letra o sílaba del nombre. Ayuda visual: mostrar imágenes relacionadas. Ayuda de terceros: un adulto puede susurrar la respuesta correcta al oído."
    },
    26: {
        "Nombre": "Selección/preferencia '¿quieres?'",
        "Objetivo": "Responder verbalmente con 'sí' o 'no' según sus gustos y deseos.",
        "Procedimiento": "Sentar al alumno(a) y mostrarle un objeto de su interés. Preguntar ‘¿Quieres?’. Reforzar si responde correctamente.",
        "Ayudas": "Ayuda verbal: decir la palabra ‘sí’. Ayuda de terceros: el cuidador susurra la respuesta al oído."
    },
    27: {
        "Nombre": "¿Qué quieres?",
        "Objetivo": "Responder con el nombre del objeto deseado cuando se le muestran opciones.",
        "Procedimiento": "Presentar varios objetos de su interés y preguntar ‘¿Qué quieres?’. Reforzar si responde correctamente.",
        "Ayudas": "Ayuda verbal: nombrar la primera o última sílaba de la palabra. Ayuda de terceros: el cuidador susurra la respuesta al oído."
    },
    28: {
        "Nombre": "De función '¿Para qué sirve?'",
        "Objetivo": "Identificar y nombrar la función de los objetos.",
        "Procedimiento": "Mostrar un objeto (ej. vaso) y preguntar ‘¿Para qué sirve?’. El alumno(a) debe señalar y responder correctamente.",
        "Ayudas": "Ayuda visual: mímica de la acción. Ayuda verbal: nombrar la primera o última sílaba. Ayuda de terceros: susurrar la respuesta al oído."
    },
    29: {
        "Nombre": "Posesivos Mío/Tuyo",
        "Objetivo": "Responder ‘mío’ o ‘tuyo’ a la pregunta ‘¿de quién es?’.",
        "Procedimiento": "Tocar una parte del cuerpo del terapeuta (ej. cabeza) y preguntar ‘¿de quién es esta cabeza?’. Reforzar si responde correctamente.",
        "Ayudas": "Ayuda verbal: decir la primera sílaba del pronombre. Ayuda de terceros: susurrar la respuesta al oído."
    },
    30: {
        "Nombre": "Preposiciones '¿dónde?'",
        "Objetivo": "Identificar y nombrar conceptos espaciales respondiendo ‘¿dónde está?’",
        "Procedimiento": "Dar un objeto al alumno(a) y pedirle que lo coloque en un lugar específico (ej. ‘arriba de la mesa’). Preguntar ‘¿dónde está?’.",
        "Ayudas": "Ayuda física: guiar la mano del alumno(a) para colocar el objeto. Ayuda visual: señalar el lugar con el dedo. Ayuda verbal: decir la primera sílaba de la palabra."
    },
    31: {
        "Nombre": "Sensaciones ¿Dónde sientes?",
        "Objetivo": "Identificar y nombrar la parte del cuerpo en donde sienta el estímulo.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y captar su atención. Se trabaja con material sensorial táctil por ejemplo: algún objeto frío, caliente, húmedo, suave, áspero, cosquillas. Se coloca el objeto en una parte del cuerpo del alumno(a) al mismo tiempo en el que se realiza la pregunta ¿dónde sientes frío? o ¿dónde sientes húmedo? Si responde correctamente con la parte del cuerpo en donde se le colocó el objeto se refuerza. Si no hay respuesta o nombra una parte del cuerpo distinta se corrige y brindan las ayudas necesarias.",
        "Ayudas": "Ayudas visuales: señalar con dedo índice la parte del cuerpo en donde debería de sentir la sensación. Ayudas verbales: Nombrar la parte del cuerpo en donde se colocó el objeto. Nombrar la primera o última sílaba de la palabra. Decir la respuesta correcta antes de realizar la pregunta. Ayuda física: tomar la mano del alumno(a) y colocarla en donde debió de sentir la sensación."
    },
    32: {
        "Nombre": "Acciones ¿Qué está haciendo?",
        "Objetivo": "Identificar y nombrar la acción que se le muestra.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y captar su atención. Colocar tres tarjetas de acciones sobre la mesa al mismo tiempo en el que se nombran. Se le pide que entregue una acción ejemplo: “dame el que está comiendo”. Si lo hace correctamente se refuerza y se continúa con la pregunta ¿Qué está haciendo? Si responde correctamente con la acción se refuerza. Si no hay respuesta o es incorrecta se brindan las ayudas necesarias para cumplir con el objetivo.",
        "Ayudas": "Ayuda verbal: Mencionar la primera o última sílaba de la palabra o la palabra completa. Decir la respuesta correcta antes de realizar la pregunta. Ejemplo: está saltando ¿Qué está haciendo? Ayuda de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta realice la pregunta con el objetivo de que el alumno(a) la repita."
    },
    33: {
        "Nombre": "¿Qué estás haciendo?",
        "Objetivo": "Identificar y nombrar la acción que está realizando.",
        "Procedimiento": "Ponerse frente al alumno(a) y captar su atención. Realizar una acción con el alumno(a), por ejemplo, ‘vamos a bailar, vamos a saltar’. Preguntar (nombre del alumno(a) ‘¿qué estás haciendo?’ Si da la respuesta correcta como ‘saltando’ o ‘estoy saltando’ reforzar. Si no hay respuesta o es incorrecta brindar ayudas para poder lograr el objetivo.",
        "Ayudas": "Ayudas verbales: nombrar la primera o última sílaba de la palabra. Decir la respuesta correcta antes de realizar la pregunta. Ejemplo: saltando ¿qué estás haciendo? Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta realice la pregunta con el objetivo de que el alumno(a) la repita."
    },
    34: {
        "Nombre": "¿Qué estoy haciendo?",
        "Objetivo": "Identificar y nombrar la acción del terapeuta o de la persona que realice la pregunta.",
        "Procedimiento": "Colocarse frente al alumno(a) y captar su atención con instrucción ‘mira’ seguido de una acción por ejemplo: saltar. Realizar la pregunta ‘¿Qué estoy haciendo?’ Si responde correctamente con la acción se refuerza. Si no hay respuesta o es incorrecta brindar ayudas para poder lograr el objetivo.",
        "Ayudas": "Ayudas verbales: nombrar la primera o última sílaba de la palabra. Decir la respuesta correcta antes de realizar la pregunta. Ejemplo: saltando ¿qué estoy haciendo? Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta realice la pregunta con el objetivo de que el alumno(a) la repita."
    },
    35: {
        "Nombre": "Descripciones",
        "Objetivo": "Nombrar adjetivos calificativos.",
        "Procedimiento": "Utilizar el apoyo de tarjetas de objetos, de lugares y de emociones. Sentar al alumno(a) frente al terapeuta y captar su atención. Presentar una categoría y una tarjeta a la vez. Realizar la pregunta ¿cómo es? O ¿Cómo se siente? dependiendo de la tarjeta que se le haya presentado. Si la respuesta es correcta o hay una aproximación se refuerza. Si no hay respuesta o es incorrecta se le brindan ayudas.",
        "Ayudas": "Ayudas verbales: Decir la respuesta correcta ejemplo: ‘es grande y rojo’ o ‘se siente triste’ dependiendo de la imagen y realizar la pregunta nuevamente. En caso de ser dos conceptos descriptivos brindar ayuda mencionando el conector ‘y’. ejemplo: ‘___y____’. Nombrar la primera sílaba de la palabra. Ayudas de terceros: el cuidadora(a) se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta realice la pregunta como objetivo de que el alumno(a) la repita."
    },
    36: {
        "Nombre": "Explicación: ¿por qué?",
        "Objetivo": "Responder a preguntas de tipo ¿por qué? Discriminando explicando y eligiendo el estímulo que no corresponde a la categoría.",
        "Procedimiento": "INTRUSO: Sentarse frente al alumno(a) y captar su atención. Se colocan cuatro tarjetas sobre la mesa, tres de una misma categoría y una de categoría distinta. Ejemplo: 3 animales y 1 fruta. Se realiza la pregunta ‘¿Cuál no va?’ si responde correctamente se refuerza y se realiza la pregunta ‘¿Por qué?’. Si responde correctamente ‘porque es una fruta y no un animal’ o alguna aproximación se refuerza. Si no hay respuesta o es incorrecta se le brindan ayudas. DE ACCIÓN: Presentar las acciones: el niño está comiendo porque tiene hambre ¿Por qué está comiendo? Si el alumno no responde brinde ayudas verbales: ‘Porque tiene hambre’",
        "Ayudas": "Ayuda visual: Señalar la tarjeta que no va. Ayudas verbales: Mencionar la respuesta correcta. Brindar ayudas parciales para formar la oración. Ejemplo: Terapeuta: porque es una____ y no es un_____. Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta realice la pregunta como objetivo de que el alumno(a) la repita."
    },
    37: {
        "Nombre": "Auto cíclicas",
        "Objetivo": "Realizar oraciones cortas de 3-4 palabras.",
        "Procedimiento": "Utilizar el apoyo de imágenes u objetos de animales, frutas, verduras o distintos objetos de uso común. Sentar al alumno(a) frente al terapeuta y captar su atención. El terapeuta toma un objeto o imagen y se le brinda otro al alumno(a). El terapeuta comienza nombrando lo que tiene comenzando la oración con ‘yo tengo (nombrar el objeto)’ seguido de la pregunta ‘¿y tú?’. Si lo hace correctamente se refuerza. Si no hay respuesta o es incorrecta se le brindan ayudas para poder lograr el objetivo. Nota: Al ir avanzando en el programa se va retirando la mención de la pregunta ‘¿y tú?’ para que más adelante responda sin necesidad de un apoyo auditivo.",
        "Ayudas": "Ayudas verbales: Se le brinda la respuesta correcta para que el alumno(a) lo repita. Mencionar ‘yo tengo’ con el objetivo de que el alumno(a) repita ambas palabras se agrega una ayuda visual señalando el objeto para que el alumno(a) termine la oración. Ayuda física: Tomar la mano del alumno(a) y dirigirla hacia su pecho al mismo tiempo en el que se le dice ‘yo’ haciendo énfasis para que repita. Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta diga la oración como objetivo de que el alumno(a) la repita."
        },
    38: {
        "Nombre": "Yo tengo",
        "Objetivo": "Realizar oraciones cortas de 3-4 palabras.",
        "Procedimiento": "Utilizar el apoyo de imágenes u objetos de animales, frutas, verduras o distintos objetos de uso común. Sentar al alumno(a) frente al terapeuta y captar su atención. El terapeuta toma un objeto o imagen y se le brinda otro al alumno(a). El terapeuta comienza nombrando lo que tiene comenzando la oración con ‘yo tengo (nombrar el objeto)’ seguido de la pregunta ‘¿y tú?’. Si lo hace correctamente se refuerza. Si no hay respuesta o es incorrecta se le brindan ayudas para poder lograr el objetivo. Nota: Al ir avanzando en el programa se va retirando la mención de la pregunta ‘¿y tú?’ para que más adelante responda sin necesidad de un apoyo auditivo.",
        "Ayudas": "Ayudas verbales: Se le brinda la respuesta correcta para que el alumno(a) lo repita. Mencionar ‘yo tengo’ con el objetivo de que el alumno(a) repita ambas palabras se agrega una ayuda visual señalando el objeto para que el alumno(a) termine la oración. Ayuda física: Tomar la mano del alumno(a) y dirigirla hacia su pecho al mismo tiempo en el que se le dice ‘yo’ haciendo énfasis para que repita. Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta diga la oración como objetivo de que el alumno(a) la repita."
    },
    39: {
        "Nombre": "Yo veo",
        "Objetivo": "Realizar oraciones cortas de 3-4 palabras.",
        "Procedimiento": "Se comienzan nombrando objetos que estén sobre el escritorio y conforme se avanza en el programa se alarga la distancia. Colocar un objeto sobre la mesa. Captar la atención del alumno(a). Señalar el objeto y decir la oración ‘yo veo un (nombre del objeto)’ y realizar la pregunta ‘¿y tú?’. El alumno(a) deberá de nombrar lo que está viendo y si lo realiza se refuerza. Si la respuesta es incorrecta o no la hay se le brindan ayudas. Nota: al ir avanzando en el programa se va retirando la mención de la pregunta ‘¿y tú?’ para que más adelante responda sin necesidad de un apoyo auditivo.",
        "Ayudas": "Ayudas verbales: Decir la respuesta correcta antes de realizar la pregunta para que la repita. Mencionar ‘yo veo’ con el objetivo de que el alumno(a) repita ambas palabras se agrega una ayuda visual señalando el objeto para que el alumno(a) termine la oración. Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta diga la oración como objetivo de que el alumno(a) la repita. Ayuda física: Tomar la mano del alumno(a) y dirigirla hacia su pecho al mismo tiempo en el que se le dice ‘yo’ haciendo énfasis para que repita."
    },
    40: {
        "Nombre": "Me gusta/no me gusta",
        "Objetivo": "Nombrar cosas o situaciones que sean de su gusto/agrado y de su desagrado.",
        "Procedimiento": "Sentar al alumno(a) frente al terapeuta y captar su atención. El terapeuta comienza nombrando algún objeto o situación comenzando la oración con ‘a mí me gusta… ¿y a ti?’. Se realiza una pausa para que el alumno(a) conteste con algo que sea de su agrado tomando como modelo la oración que previamente dijo el terapeuta. Si la respuesta es correcta o hay una buena aproximación se refuerza. Si no hay respuesta o es incorrecta se le brindan ayudas para poder lograr con el objetivo del programa. Nota: al ir avanzando en el programa se va retirando la mención de la pregunta ‘¿y a ti?’ para que más adelante responda sin necesidad de un apoyo auditivo.",
        "Ayudas": "Ayudas verbales: Nombrar la primera o última sílaba de las palabras. Decir partes de la oración para que el alumno(a) la complete. Decir la oración completa con el objetivo de que el alumno(a) la repita. Ayudas visuales: mostrar imágenes de objetos que sean de su agrado. Decir ‘no’ con la cabeza al mismo tiempo en el que se dice la oración ‘a mí no me gusta’. Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta diga la oración como objetivo de que el alumno(a) la repita."
    },
    41: {
        "Nombre": "Formación de oraciones",
        "Objetivo": "Desarrollar y formar oraciones compuestas por sujeto, verbo y predicado.",
        "Procedimiento": "Con apoyo de tarjetas de personas, profesiones, objetos y acciones. Presentar al alumno(a) una tarjeta de cada uno de esos campos, ejemplo: alumno(a), acción de comer y la imagen de una fruta. Colocarlas en el orden correcto y modelarle la respuesta: ‘el alumno(a) está comiendo manzana’. Se le pregunta: ¿qué está pasando? Si el alumno(a) responde correctamente o hay alguna aproximación se refuerza, de lo contrario se brindan las ayudas necesarias.",
        "Ayudas": "Ayuda visual: señalar las imágenes mientras se forma la oración. Ayuda verbal: Mencionar la oración por partes para que el alumno(a) vaya repitiendo. Decir partes de la oración dejando espacios para que el alumno(a) la complete."
    },
    42: {
        "Nombre": "Experiencias",
        "Objetivo": "Narrar y expresar experiencias y sucesos que hayan ocurrido durante el día o en días pasados.",
        "Procedimiento": "Es importante preguntar con anterioridad al padre o tutor tres actividades que se realizaron en ese día. Sentarse frente al alumno(a) y captar su atención. El terapeuta cuenta algún suceso de su día ejemplo: ‘Yo en la mañana desayuné huevo’ y se realiza la pregunta ‘¿y tú?’. Si hay respuesta o aproximación se refuerza. Si no hay respuesta se brindan ayudas. Nota: al ir avanzando en el programa se va retirando la mención de la pregunta ‘¿y tú?’ o de preguntas relacionadas a sucesos para que más adelante responda sin necesidad de un apoyo auditivo. Nota: se comienza con sucesos del día y conforme avanza con sucesos de días pasados.",
        "Ayudas": "Ayudas verbales: Decir partes de la oración para que el alumno(a) la complete. Decir la oración completa con el objetivo de que el alumno(a) la repita. Ayudas de terceros: el cuidador se coloca a un lado del alumno(a) y le dice la respuesta correcta al oído inmediatamente después de que el terapeuta diga la oración como objetivo de que el alumno(a) la repita. Ayudas visuales: Presentarles imágenes relacionadas a los sucesos relacionados con las preguntas."
    },
    43: {
        "Nombre": "Correcta articulación de fonemas",
        "Objetivo": "Que el alumno(a) emita correctamente los sonidos de los fonemas comenzando de manera aislada, después en sílabas, palabras y oraciones cortas.",
        "Procedimiento": "Sentarse frente al alumno(a) y captar su atención. Realizar los distintos movimientos de boca y lengua dependiendo del fonema que se está trabajando. Si hay respuesta o aproximación se refuerza. Si no hay respuesta se brindan ayudas. Continuar agregando la vocal ‘a’. Trabajar dicho fonema en cadenas silábicas integrando todas las vocales. Practicar el fonema deseado en palabras cortas y por último complejas. Si hay respuesta o aproximación se refuerza. Si no hay respuesta se brindan ayudas. Nota: realizar esta actividad preferentemente frente a un espejo.",
        "Ayudas": "Ayudas verbales: cambiar entonación, alargar sonido de vocales."
    },
    44: {
        "Nombre": "Atención Focal Visual Auditiva",
        "Objetivo": "Desarrollar la capacidad de seleccionar y centrar la atención en un estímulo en particular.",
        "Procedimiento": "Seguimiento de estímulo visual. Atender un estímulo auditivo: ‘alto’. Seguir el ritmo: replicar una secuencia rítmica. Replicar una secuencia visual.",
        "Ayudas": "Ayuda visual: Uso de colores brillantes para estímulos visuales. Ayuda auditiva: Uso de tonos altos o diferenciados."
    },
    45: {
        "Nombre": "Atención Sostenida Visual Auditiva",
        "Objetivo": "Mantener la atención durante un periodo prolongado de tiempo mientras se lleva a cabo una determinada tarea (interviene la motivación/fatiga).",
        "Procedimiento": "Lotería de sonidos: cuadernillo de actividades cognitivas y cuadernillo de ejercicio. Discriminación visual. Cuadernillo de ejercicios. Encontrar las diferencias. Sopa de letras. La lotería. Replicar modelos visuales.",
        "Ayudas": "Ayuda visual: Uso de marcadores o señalizadores para mantener el enfoque. Ayuda verbal: Recordatorios periódicos."
    },
    46: {
        "Nombre": "Atención Dividida",
        "Objetivo": "Desarrollar la capacidad de atender a más de un estímulo a la vez.",
        "Procedimiento": "La lotería. La pelotita: actividad motora y de fluidez verbal. Cuadernillo de actividades. Secuencia de colores y ritmos: cuadernillo de actividades.",
        "Ayudas": "Ayuda visual: Colores contrastantes. Ayuda motora: Guiar físicamente en actividades de movimiento."
    },
    47: {
        "Nombre": "Atención Conjunta",
        "Objetivo": "Estimular y desarrollar la habilidad de compartir intereses por un mismo objeto o actividad.",
        "Procedimiento": "Pasar la pelota: con ambas manos, lanzando, patearla con el pie. Cuadernillo de actividades. Juegos de mesa: memorama, dominó, el uno. Juegos sensoriales (cosquillas, dar vueltas, etc.). Juego de marionetas. Burbujas. Juegos colaborativos.",
        "Ayudas": "Ayuda visual: Demostraciones explícitas. Ayuda social: Instrucciones claras y modelado de comportamientos."
    },
    48: {
        "Nombre": "Memoria de trabajo",
        "Objetivo": "Memorizar y recuperar información para ejecutar una tarea a corto plazo a partir de estímulos visuales o auditivos.",
        "Procedimiento": "Recuperación visual y auditiva: cuadernillo de actividades. Replicar una secuencia visual. Secuencia de colores y ritmos: cuadernillo de actividades. Memorizar una canción. Memorizar un trabalenguas. Memorizar un poema.",
        "Ayudas": "Ayuda visual: Diagramas o esquemas. Ayuda auditiva: Repeticiones y ritmo."
    },
    49: {
        "Nombre": "Memoria Episódica",
        "Objetivo": "Recordar momentos, lugares y emociones de un suceso de manera estructurada.",
        "Procedimiento": "Juego de preguntas de información personal y experiencias. Ajedrez. Realizar actividades de lectura de comprensión. Autobiografía con imágenes. Contar una historia y pedirle que la relate nuevamente. Álbum de fotos.",
        "Ayudas": "Ayuda visual: Notas adhesivas, esquemas temporales. Ayuda verbal: Repeticiones y cuestionarios."
    },
    50: {
        "Nombre": "Control inhibitorio",
        "Objetivo": "Regular y concientizar las respuestas automatizadas.",
        "Procedimiento": "Actividades go-no go, mencionar el artículo contrario o diferente al que se muestra: Con figuras, Colores, Movimientos motores, Cuadernillo de actividades.",
        "Ayudas": "Ayuda visual: Álbum de fotos, mapas conceptuales. Ayuda verbal: Preguntas guiadas y recordatorios."
    },
    51: {
        "Nombre": "Flexibilidad",
        "Objetivo": "Trabajar en la adaptabilidad y regulación ante cambios inesperados proponiendo diferentes actividades.",
        "Procedimiento": "Ir del punto A al punto B: Cuadernillo de actividades. Personajes revueltos: cuadernillo de actividades y de ejercicios. Colores locos: cuadernillo de actividades. Historias absurdas: cuadernillo de actividades. Solución de problemas: cuadernillo de actividades.",
        "Ayudas": "Ayuda visual: Señales de alto y sigue. Ayuda verbal: Instrucciones claras y refuerzo positivo."
    },
    52: {
        "Nombre": "Semejanzas",
        "Objetivo": "Desarrollar el pensamiento abstracto realizando inferencias básicas, relacionando conceptos de diferentes categorías.",
        "Procedimiento": "Tarjetas: se colocan dos imágenes y se hace la pregunta ‘¿en qué se parecen?’ más adelante se retiran las ayudas visuales. Cuadernillo de actividades.",
        "Ayudas": "Ayuda visual: Diagramas de flujo. Ayuda verbal: Explicaciones de posibles escenarios."
    },
    53: {
        "Nombre": "Diferencias",
        "Objetivo": "Desarrollar el pensamiento abstracto identificando diferencias entre dos conceptos.",
        "Procedimiento": "Tarjetas: se colocan dos imágenes y se realiza la pregunta ‘¿en qué son diferentes?’ más adelante se retiran las ayudas visuales.",
        "Ayudas": "Ayuda visual: Tarjetas con pistas visuales. Ayuda verbal: Preguntas guiadas."
    },
    54: {
        "Nombre": "Fluidez Verbal",
        "Objetivo": "Incrementar la comunicación, el vocabulario y la fluidez.",
        "Procedimiento": "El basta: cuadernillo de actividades. Mencionar cierta cantidad de artículos de una categoría: cuadernillo de actividades. Palabras por minuto: cuadernillo de actividades.",
        "Ayudas": "Ayuda visual: Tarjetas con pistas visuales. Ayuda verbal: Preguntas guiadas."
    },
    55: {
        "Nombre": "Temporalidad",
        "Objetivo": "Comprender sentidos temporales: ayer, hoy, mañana.",
        "Procedimiento": "Primero y después: cuadernillo de actividades. Juego de preguntas: cuadernillo de actividades.",
        "Ayudas": "Ayuda verbal: Proveer ejemplos y listas de palabras. Ayuda visual: Tablas de categorías."
    },
    56: {
        "Nombre": "Adivinanzas",
        "Objetivo": "Estimular el pensamiento abstracto a partir de conceptos descriptivos.",
        "Procedimiento": "Adivinanzas: cuadernillo de ejercicios.",
        "Ayudas": "Ayuda visual: Calendarios y líneas de tiempo. Ayuda verbal: Explicaciones de secuencias temporales."
    },
    57: {
        "Nombre": "Verbos Mentales",
        "Objetivo": "Desarrollar el pensamiento abstracto, juicio y análisis por medio de conceptos abstractos, identificando deseos, pensamientos, gustos y creencias ajenas a las de uno mismo.",
        "Procedimiento": "Creo, pienso, sueño, deseo, imagino, siento, mentiras, inferencias: cuadernillo de actividades ‘verbos mentales’. Nota: para la elaboración del plan de trabajo se mencionan los verbos a trabajar y se desarrolla la actividad.",
        "Ayudas": "Ayuda verbal: Dar pistas y simplificar las adivinanzas."
    },
    58: {
        "Nombre": "Uso de metáforas",
        "Objetivo": "Identificar y comprender el sentido figurado de las oraciones.",
        "Procedimiento": "Metáforas: ‘frases de sentido figurado’. Cuadernillo de ejercicios.",
        "Ayudas": "Ayuda verbal: Ejemplos y escenarios prácticos."
    },
    59: {
        "Nombre": "Absurdos",
        "Objetivo": "Desarrollar el pensamiento lógico proponiendo situaciones para que el alumno discrimine si un escenario es posible o absurdo.",
        "Procedimiento": "Absurdos: cuadernillo de actividades.",
        "Ayudas": "Ayuda verbal: Ejemplos y explicaciones de metáforas."
    },
    60: {
        "Nombre": "Planificación",
        "Objetivo": "Incrementar la organización y planificación.",
        "Procedimiento": "Siguiendo el caminito: cuadernillo de ejercicios. Laberintos. Juegos de mesa.",
        "Ayudas": "Ayuda visual: Ilustraciones de escenarios. Ayuda verbal: Explicaciones y preguntas dirigidas."
    },
    61: {
        "Nombre": "Motricidad fina",
        "Objetivo": "Desarrollar habilidades motoras finas para el agarre y escritura.",
        "Procedimiento": "Colocar pinzas de ganchos en un recipiente. Ensartar cuentas. Pinchar los bordes de una figura con un picadientes sobre una superficie de hielo seco. Realizar figuras con plastilina. ‘Pesca’: atrapar tapaderas de plástico colocadas en un recipiente con agua con palillos.",
        "Ayudas": "Ayuda visual: Tablas y esquemas de planificación. Ayuda verbal: Instrucciones detalladas."
    },
    62: {
        "Nombre": "Conducta adaptativa",
        "Objetivo": "Regular y adecuar la conducta en los diferentes contextos sociales, desarrollando las habilidades sociales y prácticas que permitan responder a las circunstancias de la vida.",
        "Procedimiento": "Público vs privado. Historias sociales. Cuadernillo de actividades.",
        "Ayudas": "Ayuda visual: Demostraciones explícitas. Ayuda física: Guiar los movimientos del niño."
    },
    63: {
        "Nombre": "Tolerancia a la frustración",
        "Objetivo": "Incrementar los tiempos de tolerancia y espera ante actividades u objetos deseados, desarrollando las habilidades de regulación de impulsos de frustración, aumentando la paciencia y tolerancia ante situaciones desagradables.",
        "Procedimiento": "Programa de ‘espera’ del cuadernillo de actividades. Historias sociales. Ejercicios de respiración. Identificar en historias sociales situaciones de frustración y brindar actividades o herramientas para la situación en particular.",
        "Ayudas": "Ayuda verbal: Explicaciones de contextos y conductas apropiadas."
    },
    64: {
        "Nombre": "Identificación de emociones",
        "Objetivo": "Identificar y reconocer las emociones.",
        "Procedimiento": "Historias sociales: Identificar situaciones que me provoquen felicidad, enojo, tristeza, disgusto. Dibujar o pintar cómo me hace sentir cada emoción. Describir verbalmente o por escrito las emociones. Se explica la importancia de comprender las emociones.",
        "Ayudas": "Ayuda verbal: Estrategias de afrontamiento. Ayuda visual: Diagramas de emociones y soluciones."
    },
    65: {
        "Nombre": "Gestión de emociones",
        "Objetivo": "Identificar, modular y canalizar las emociones.",
        "Procedimiento": "Historias sociales: identificar las emociones que me genera una situación en particular y realizar una lista de actividades que el paciente pueda hacer en cada una de esas situaciones (la importancia de la respiración y de expresar mi sentir de manera amable).",
        "Ayudas": "Ayuda visual: Caras de emociones. Ayuda verbal: Explicaciones de emociones y ejemplos."
    },
    66: {
        "Nombre": "Autoconocimiento",
        "Objetivo": "Desarrollar autoconciencia de uno mismo, reconocer fortalezas, áreas de oportunidad, gustos, disgustos, entre otros.",
        "Procedimiento": "Realizar por escrito o con recortes una lista de cosas que le gustan y de cosas que no le gustan al paciente. Realizar una lista con recortes o escrito de fortalezas y áreas de oportunidad, analizarlo y platicarlo.",
        "Ayudas": "Ayuda verbal: Técnicas de regulación emocional. Ayuda visual: Diagramas de emociones."
    }
        }

# Verificar si el archivo 'programas.json' existe
if os.path.exists("programas.json"):
    # Si el archivo existe, cargar su contenido en programas
    with open("programas.json", "r") as file:
        programas_json = json.load(file)
    
    # Limpiar y actualizar el diccionario `programas` con los datos del archivo
    programas.clear()
    programas.update({int(k): v for k, v in programas_json.items()})
else:
    # Si el archivo no existe, crearlo con un diccionario vacío
    with open("programas.json", "w") as file:
        json.dump(programas, file, indent=4)


# Función para seleccionar programas
def seleccionar_programas_plan(entry_programas_seleccionados):
    # Crear una nueva ventana para seleccionar programas
    ventana_seleccion = tk.Toplevel()
    ventana_seleccion.title("Seleccionar Programas")
    ventana_seleccion.geometry("400x300")  # Tamaño de la ventana para que sea más visible

    # Frame contenedor con scrollbar
    frame_contenedor = tk.Frame(ventana_seleccion)
    frame_contenedor.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

    # Configuración del grid de la ventana
    ventana_seleccion.grid_rowconfigure(0, weight=1)
    ventana_seleccion.grid_columnconfigure(0, weight=1)

    # Canvas y scrollbar
    canvas = tk.Canvas(frame_contenedor)
    scrollbar = tk.Scrollbar(frame_contenedor, orient="vertical", command=canvas.yview)
    frame_scrollable = tk.Frame(canvas)

    # Configuración de scrollbar
    frame_scrollable.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    # Colocar el frame scrollable en el canvas
    canvas.create_window((0, 0), window=frame_scrollable, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Colocar canvas y scrollbar en el grid
    canvas.grid(row=0, column=0, sticky="nsew")
    scrollbar.grid(row=0, column=1, sticky="ns")

    # Expandir el frame contenedor y el canvas para ajustarse al tamaño de la ventana
    frame_contenedor.grid_rowconfigure(0, weight=1)
    frame_contenedor.grid_columnconfigure(0, weight=1)

    seleccionados = []  # Lista para almacenar los programas seleccionados

    for i, (programa_id, programa_info) in enumerate(programas.items()):
        var = tk.IntVar()  # Variable que se usará para marcar el checkbox
        check = tk.Checkbutton(frame_scrollable, text=programa_info["Nombre"], variable=var)
        check.grid(row=i, column=0, sticky='w', pady=2)  # Usar grid para colocar los checkboxes
        seleccionados.append((programa_id, var))  # Guardar el id y el estado de selección

    # Función que se ejecuta cuando se confirman los programas seleccionados
    def confirmar_seleccion():
        seleccion_final = [str(id) for id, var in seleccionados if var.get() == 1]

        # Actualizar el campo "Programas Seleccionados" en la entrada principal
        entry_programas_seleccionados.delete(0, tk.END)  # Limpiar la entrada
        entry_programas_seleccionados.insert(0, ','.join(seleccion_final))  # Insertar la selección

        # Cerrar la ventana de selección
        ventana_seleccion.destroy()

    # Botón para confirmar la selección
    boton_confirmar = tk.Button(ventana_seleccion, text="Confirmar", command=confirmar_seleccion)
    boton_confirmar.grid(row=1, column=0, pady=10)  # Usar grid para colocar el botón




# Función para crear el documento de Word
def crear_documento(datos):
    import json

    # Leer los datos del archivo configuración
    with open('configuracion.json', 'r') as file:
        configuracion = json.load(file)

    # Extraer nombre y cédula
    nombre_terapeuta = configuracion["nombre_terapeuta"]
    cedula_profesional = configuracion["cedula_profesional"]
    
    doc = Document()
    # Configurar la fuente del documento a Arial 11
    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Arial'
    fuente.size = Pt(11)

    # Añadir una imagen al header del documento
    section = doc.sections[0]
    header = section.header

    # Insertar la imagen en el encabezado
    image_path_header = resource_path("Cabezera.png")
    header_paragraph = header.paragraphs[0]
    header_paragraph.add_run().add_picture(image_path_header, width=Inches(6.1))  # Ajusta el tamaño de la imagen

    # Añadir una imagen al footer del documento
    section = doc.sections[0]
    footer = section.footer

    # Insertar la imagen en el pie de página
    image_path_footer = resource_path("Fondo.png")
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.add_run().add_picture(image_path_footer, width=Inches(6.1))  # Ajusta el tamaño de la imagen
    
    # Añadir el contenido


    # Función para eliminar los bordes de una celda específica
    def remove_cell_borders(cell, borders_to_remove):
        """
        Elimina los bordes especificados de una celda.
    
        :param cell: La celda a la que se le eliminarán los bordes.
        :param borders_to_remove: Una lista de los bordes a eliminar. Ej: ['top', 'left', 'bottom', 'right']
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
    
        # Crear el elemento w:tcBorders si no existe
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
    
        # Eliminar los bordes especificados
        for border in borders_to_remove:
            border_element = tcBorders.find(qn(f'w:{border}'))
            if border_element is None:
                border_element = OxmlElement(f'w:{border}')
                tcBorders.append(border_element)
            border_element.set(qn('w:val'), 'nil')  # Eliminar el borde

    # Añadir la tabla con 3 filas y 2 columnas
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Table Grid'


    # Quitar los bordes de las celdas originales antes de combinar
    remove_cell_borders(table1.cell(0, 0), ['top', 'left'])  # Borrar bordes superior e izquierdo de la celda (0, 0)
    remove_cell_borders(table1.cell(1, 0), ['left'])  # Borrar borde izquierdo de la celda (1, 0)
    remove_cell_borders(table1.cell(2, 0), ['left', 'bottom'])  # Borrar bordes izquierdo e inferior de la celda (2, 0)

    # Combinar las celdas de la primera columna (de las 3 filas)
    a = table1.cell(0, 0)  # Primera celda de la primera columna
    b = table1.cell(2, 0)  # Última celda de la primera columna
    merged_cell = a.merge(b)  # Combinar las celdas de la primera columna y asignar a merged_cell

    # Insertar la imagen en la primera columna (ocupando las 3 filas de la primera columna)
    image_path = resource_path("Tabla.png")
    merged_cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2.5))

    # Llenar los nombres de los datos en la segunda columna y los valores en la tercera columna

    # Formateo de la celda (0, 1)
    cell_0_1 = table1.cell(0, 1).paragraphs[0]
    run_0_1 = cell_0_1.add_run('Fecha: ')
    run_0_1.bold = True
    cell_0_1.add_run(datos["Fecha"])

    # Formateo de la celda (1, 1)
    cell_1_1 = table1.cell(1, 1).paragraphs[0]
    run_1_1 = cell_1_1.add_run('Número de Expediente: ')
    run_1_1.bold = True
    cell_1_1.add_run(datos["Número de Expediente"])

    # Formateo de la celda (2, 1)
    cell_2_1 = table1.cell(2, 1).paragraphs[0]
    run_2_1 = cell_2_1.add_run('Unidad: ')
    run_2_1.bold = True
    cell_2_1.add_run(datos["Unidad"])
 
    # Añadir una tabla con 1 fila y 1 columna
    table6 = doc.add_table(rows=1, cols=1)
    cell = table6.cell(0, 0)

    # Cambiar el color de fondo de la celda a guinda
    tcPr = cell._element.get_or_add_tcPr()  # Obtener las propiedades de la celda
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '800000')  # Código hexadecimal para guinda (color de fondo)
    tcPr.append(shading_elm)

    # Añadir texto a la celda
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

    # Añadir el texto "PLAN DE INTERVENCIÓN PSICOSOCIAL"
    run = paragraph.add_run("PLAN DE INTERVENCIÓN PSICOSOCIAL")
    run.bold = True
    run.font.size = Pt(11)  # Tamaño de la fuente
    run.font.color.rgb = RGBColor(255, 255, 255)  # Color del texto (blanco)

    # Espacio entre tablas
    doc.add_paragraph('')

    # Segunda tabla: Encabezados en la primera columna y valores en la segunda columna
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'

    # Encabezados en la primera columna

    # Formateo de la celda (0, 0)
    cell = table2.cell(0, 0).paragraphs[0]
    run = cell.add_run('Nombre: ')
    run.bold = True
    cell.add_run(datos["Nombre del Paciente"])

    # Formateo de la celda (0, 1)
    cell = table2.cell(0, 1).paragraphs[0]
    run = cell.add_run('Edad: ')
    run.bold = True
    cell.add_run(datos["Edad"])

    # Formateo de la celda (1, 0)
    cell = table2.cell(1, 0).paragraphs[0]
    run = cell.add_run('Fecha de Nacimiento: ')
    run.bold = True
    cell.add_run(datos["Fecha de Nacimiento"])

    # Formateo de la celda (1, 1)
    cell = table2.cell(1, 1).paragraphs[0]
    run = cell.add_run('Diagnóstico: ')
    run.bold = True
    cell.add_run(datos["Diagnóstico"])

    # Formateo de la celda (2, 1)
    cell = table2.cell(2, 1).paragraphs[0]
    run = cell.add_run('Num. de Sesiones por semana: ')
    run.bold = True
    cell.add_run(datos["Sesiones por Semana"])

    # Formateo de la celda (2, 0)
    cell = table2.cell(2, 0).paragraphs[0]
    run = cell.add_run('Intervención en terapia: ')
    run.bold = True
    cell.add_run(datos["Área de Intervención"])
    
    # Espacio entre tablas
    doc.add_paragraph('')

    # Añadir una tabla con 1 fila y 1 columna
    table7 = doc.add_table(rows=1, cols=1)
    cell = table7.cell(0, 0)

    # Cambiar el color de fondo de la celda a guinda
    tcPr = cell._element.get_or_add_tcPr()  # Obtener las propiedades de la celda
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '800000')  # Código hexadecimal para guinda (color de fondo)
    tcPr.append(shading_elm)

    # Añadir texto a la celda
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

    # Añadir el texto "REFORZADORES"
    run = paragraph.add_run("REFORZADORES")
    run.bold = True
    run.font.size = Pt(11)  # Tamaño de la fuente
    run.font.color.rgb = RGBColor(255, 255, 255)  # Color del texto (blanco)



    # Tercera tabla: Encabezados en la primera columna y valores en la segunda columna
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Table Grid'
    
    # Encabezados en la primera columna
    hdr_cells1 = table1.rows[0].cells
    hdr_cells1[0].text = 'Comestibles'
    hdr_cells1[1].text = datos['Reforzadores Comestibles']

    row_cells1 = table1.rows[1].cells
    row_cells1[0].text = 'Tangibles'
    row_cells1[1].text = datos['Reforzadores Tangibles']

    row_cells2 = table1.rows[2].cells
    row_cells2[0].text = 'Sociales'
    row_cells2[1].text = datos['Reforzadores Sociales']

    # Espacio entre tablas
    doc.add_paragraph('')

    # Crear una tabla de una celda para las observaciones clínicas
    table9 = doc.add_table(rows=1, cols=1)
    table9.style = 'Table Grid'

    # Eliminar los bordes de la tabla
    def remove_borders(table):
        tbl = table._element
        tblPr = tbl.xpath('.//w:tblPr')[0]  # Buscar el elemento tblPr
        tblBorders = OxmlElement('w:tblBorders')
    
        # Crear los bordes vacíos
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')  # Configura el valor a 'none' para eliminar bordes
            tblBorders.append(border)
    
        tblPr.append(tblBorders)  # Agregar la configuración de bordes a la tabla

    # Insertar las observaciones clínicas en la tabla
    cell = table9.cell(0, 0)
    cell.text = ('El plan de intervención en terapia integral tiene como objetivo estimular habilidades cognitivas '
             'relacionadas con los procesos mentales implicados en la ejecución y planeación de tareas o actividades, '
             'permitiendo adaptar la conducta ante los diferentes contextos sociales. Así como el desarrollo de habilidades '
             'relacionadas con el lenguaje, la comprensión y expresión de ideas. Varía según el perfil del alumno.')

    # Llamar la función para eliminar los bordes de la tabla
    remove_borders(table9)

    # Espacio entre tablas
    doc.add_paragraph('')
    
    # Sección de observaciones clínicas
    parrafo = doc.add_paragraph()
    run = parrafo.add_run('Observaciones Clínicas')
    run.bold = True  # Poner en negritas
    run.underline = True  # Subrayar

    # Crear una tabla de una celda para las observaciones clínicas
    table4 = doc.add_table(rows=1, cols=1)
    table4.style = 'Table Grid'

    # Insertar las observaciones clínicas en la tabla
    cell = table4.cell(0, 0)
    cell.text = datos['Observaciones Clínicas']

    # Insertar un salto de página
    doc.add_page_break()

    # Añadir una tabla con 1 fila y 1 columna
    table8 = doc.add_table(rows=1, cols=1)
    cell = table8.cell(0, 0)

    # Cambiar el color de fondo de la celda a guinda
    tcPr = cell._element.get_or_add_tcPr()  # Obtener las propiedades de la celda
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '800000')  # Código hexadecimal para guinda (color de fondo)
    tcPr.append(shading_elm)

    # Añadir texto a la celda
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

    # Añadir el texto "Plan de Trabajo"
    run = paragraph.add_run("Plan de Trabajo")
    run.bold = True
    run.font.size = Pt(11)  # Tamaño de la fuente
    run.font.color.rgb = RGBColor(255, 255, 255)  # Color del texto (blanco)

    
    # Crear la tabla con borde continuo
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = 'Table Grid'

    # Aplicar bordes a las celdas
    for fila in tabla.rows:
        for celda in fila.cells:
            celda._element.get_or_add_tcPr().append(parse_xml(
                r'<w:tcBorders {}>'
                r'<w:top w:val="single" w:sz="4"/>'
                r'<w:left w:val="single" w:sz="4"/>'
                r'<w:bottom w:val="single" w:sz="4"/>'
                r'<w:right w:val="single" w:sz="4"/>'
                r'</w:tcBorders>'.format(nsdecls('w'))))

    # Añadir encabezados a la tabla
    encabezado = tabla.rows[0].cells
    encabezado[0].text = 'Nombre'
    encabezado[1].text = 'Objetivo'
    encabezado[2].text = 'Procedimiento'
    encabezado[3].text = 'Ayudas'

    # Aplicar negritas a los textos de los encabezados
    for celda in encabezado:
        for parrafo in celda.paragraphs:
            for run in parrafo.runs:
                run.bold = True

    # Añadir los programas seleccionados
    for programa_id in datos["Programas Seleccionados"]:
        # Convertir programa_id a entero para asegurarse de que coincida con las claves del diccionario
        programa_id = int(programa_id)

        # Verificar si el programa_id existe en el diccionario
        if programa_id not in programas:
            print(f"Error: El programa_id '{programa_id}' no se encuentra en el diccionario.")
            continue

        programa = programas[programa_id]
        fila = tabla.add_row().cells
        nombre_parrafo = fila[0].paragraphs[0]
        run_nombre = nombre_parrafo.add_run(programa["Nombre"])
        run_nombre.bold = True
        fila[1].text = programa["Objetivo"]
        fila[2].text = programa["Procedimiento"]
        fila[3].text = programa["Ayudas"]


    # Añadir filas al final de la tabla con Nombre y Cédula
    fila_nombre = tabla.add_row().cells
    fila_cedula = tabla.add_row().cells

    # Añadir texto centrado y formateado para Nombre
    fila_nombre[0].merge(fila_nombre[3])  # Combinar todas las celdas de la fila
    parrafo_nombre = fila_nombre[0].paragraphs[0]
    parrafo_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto
    run_nombre = parrafo_nombre.add_run(f"Nombre: {configuracion['nombre_terapeuta']}")
    run_nombre.bold = True
    run_nombre.font.size = Pt(11)  # Tamaño de fuente

    # Añadir texto centrado y formateado para Cédula
    fila_cedula[0].merge(fila_cedula[3])  # Combinar todas las celdas de la fila
    parrafo_cedula = fila_cedula[0].paragraphs[0]
    parrafo_cedula.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto
    run_cedula = parrafo_cedula.add_run(f"Cédula: {configuracion['cedula_profesional']}")
    run_cedula.bold = True
    run_cedula.font.size = Pt(11)  # Tamaño de fuente

    # Aplicar bordes continuos a las celdas combinadas
    for fila in [fila_nombre, fila_cedula]:
        celda = fila[0]
        celda._element.get_or_add_tcPr().append(parse_xml(
            r'<w:tcBorders {}>'
            r'<w:top w:val="single" w:sz="4"/>'  # Borde superior
            r'<w:left w:val="single" w:sz="4"/>'  # Borde izquierdo
            r'<w:bottom w:val="single" w:sz="4"/>'  # Borde inferior
            r'<w:right w:val="single" w:sz="4"/>'  # Borde derecho
            r'</w:tcBorders>'.format(nsdecls('w'))
        ))

    # Guardar el documento
    doc.save(f'plan_de_trabajo_{datos["Número de Expediente"]}.docx')

# Función para guardar la información en un archivo JSON
def guardar_informacion(expediente, datos):
    archivo_planes = 'planes_trabajo.json'
    
    if os.path.exists(archivo_planes):
        with open(archivo_planes, 'r') as file:
            informacion = json.load(file)
    else:
        informacion = {}
    
    # Actualizar los datos de acuerdo al número de expediente
    informacion[expediente] = datos  # Aquí asignamos los datos bajo el número de expediente
    
    with open(archivo_planes, 'w') as file:
        json.dump(informacion, file, indent=4)
    print(f"Información del expediente {expediente} guardada correctamente.")
    messagebox.showinfo("Guardado", f"Los cambios en el expediente {expediente} se han guardado correctamente.")



def recuperar_datos(expediente):
    # Cargar los datos desde el archivo JSON
    with open("planes_trabajo.json", "r") as file:
        planes_trabajo = json.load(file)
    
    # Verificar si el expediente existe
    if expediente in planes_trabajo:
        return planes_trabajo[expediente]
    else:
        raise ValueError(f"No se encontró el expediente {expediente}.")
def usar_crear_documento(expediente):
    try:
        datos = recuperar_datos(expediente)
        crear_documento(datos)  # Crear el documento con los datos recuperados
        messagebox.showinfo("Éxito", f"Documento creado para el expediente {expediente}.")
    except ValueError as e:
        messagebox.showerror("Error", str(e))        



# Función para guardar los planes de trabajo modificados
def guardar_planes_trabajo(planes_trabajo):
    archivo_planes = "planes_trabajo.json"
    with open(archivo_planes, "w") as f:
        json.dump(planes_trabajo, f, indent=4)




    
# Función para cargar las sesiones de un expediente
def cargar_sesiones(expediente_num):
    if os.path.exists('sesiones.json'):
        with open('sesiones.json', 'r') as archivo:
            sesiones = json.load(archivo)
        return sesiones.get(expediente_num, [])
    else:
        return []



# Función para calcular la edad en años y meses a partir de la fecha de nacimiento
def calcular_edad(fecha_nacimiento):
    formatos = ['%d/%m/%Y', '%Y-%m-%d']  # Agrega los formatos posibles
    for formato in formatos:
        try:
            nacimiento = datetime.strptime(fecha_nacimiento, formato)
            hoy = datetime.now()
            edad = hoy.year - nacimiento.year - ((hoy.month, hoy.day) < (nacimiento.month, nacimiento.day))
            return edad
        except ValueError:
            continue  # Intenta con el siguiente formato

    raise ValueError("El formato de la fecha no es válido.")


# Función para cargar la configuración actual
def cargar_configuracion():
    if os.path.exists('configuracion.json'):
        with open('configuracion.json', 'r') as archivo:
            configuracion = json.load(archivo)
        return configuracion
    else:
        # Si no existe el archivo de configuración, devolver un diccionario vacío
        return {"unidad": "", "nombre_terapeuta": "", "cedula_profesional": ""}

# Función para guardar los datos de configuración
def guardar_configuracion(unidad, nombre_terapeuta, cedula_profesional):
    configuracion = {
        "unidad": unidad,
        "nombre_terapeuta": nombre_terapeuta,
        "cedula_profesional": cedula_profesional
    }
    with open('configuracion.json', 'w') as archivo:
        json.dump(configuracion, archivo)
    messagebox.showinfo("Configuración", "Datos guardados correctamente.")

class ModernMainMenu:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Sistema de Planes de Trabajo")
        self.root.geometry("1024x600")
        self.root.minsize(800, 400)
        self.root.configure(bg="#F5F5DC")

        # Estilos personalizados
        self.style = ttk.Style(self.root)
        self.style.theme_use("clam")
        self.style.configure(
            "TButton",
            font=("Arial", 12),
            padding=10,
            relief="flat",
            background="#800000",  # Guinda oscuro
            foreground="white"
        )
        self.style.map(
            "TButton",
            background=[("active", "#A52A2A")],  # Guinda más claro al interactuar
            foreground=[("active", "white")]
        )

        # Dividir ventana principal en menú lateral y cuerpo principal
        self.menu_frame = tk.Frame(self.root, bg="#800000", width=250)
        self.menu_frame.pack(side="left", fill="y")

        self.body_frame = tk.Frame(self.root, bg="#FFFFFF")
        self.body_frame.pack(side="right", expand=True, fill="both")

        # Encabezado del menú lateral
        header_label = tk.Label(
            self.menu_frame,
            text="Menú",
            font=("Arial", 18, "bold"),
            bg="#800000",
            fg="white"
        )
        header_label.pack(fill="x", pady=(10, 20))

        # Botones del menú lateral
        self.create_menu_button("Crear Plan", self.crear_plan_trabajo)
        self.create_menu_button("Revisar Planes", self.revisar_planes_trabajo)
        self.create_menu_button("Directorio", self.abrir_directorio_programas)
        self.create_menu_button("Configuración", self.abrir_configuracion)
        self.create_menu_button("Informe de trabajo", self.informe_de_trabajo)


        # Área inicial en el cuerpo principal
        self.body_label = tk.Label(
            self.body_frame,
            text="Bienvenido al Sistema de Planes de Trabajo",
            font=("Arial", 16),
            bg="#FFFFFF",
            fg="#000000"
        )
        self.body_label.pack(expand=True)

    def create_menu_button(self, text, command):
        button = tk.Button(
            self.menu_frame,
            text=text,
            command=command,  # Llamar directamente a la función
            font=("Arial", 12),
            bg="#A52A2A",
            fg="white",
            relief="flat",
            activebackground="#800000",
            activeforeground="white"
        )
        button.pack(fill="x", pady=5, padx=10)
        
        
    

    def clear_body_frame(self):
        """Limpia el contenido actual del área del cuerpo principal."""
        for widget in self.body_frame.winfo_children():
            widget.destroy()
    def cargar_planes_informes():
        archivo_planes = "planes_trabajo.json"  # Usamos el archivo centralizado
    
        if not os.path.exists(archivo_planes):
            messagebox.showinfo("Error", "No se han encontrado planes de trabajo guardados.")
            return
    
        with open(archivo_planes, "r") as f:
            planes_trabajo = json.load(f)
    
        if not planes_trabajo:
            messagebox.showinfo("Error", "No hay planes de trabajo disponibles.")
            return
    
        self.clear_body_frame()
    
        # Título
        tk.Label(
            self.body_frame,
            text="Revisar planes de trabajo",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#800000"
        ).grid(row=0, column=0, pady=10)
    
        # Configurar el grid para el body_frame
        self.body_frame.grid_rowconfigure(1, weight=1)
        self.body_frame.grid_columnconfigure(0, weight=1)
    
        # Contenedor principal
        revisar_frame = tk.Frame(self.body_frame, bg="#FFFFFF")
        revisar_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
    
        # Configurar el grid para revisar_frame
        revisar_frame.grid_rowconfigure(0, weight=1)
        revisar_frame.grid_columnconfigure(0, weight=1)
    
        # Crear el Treeview
        tree = ttk.Treeview(revisar_frame, columns=("Expediente"), show="headings", height=10)
        tree.heading("Expediente", text="Expediente")
        tree.column("Expediente", anchor="w")
        tree.grid(row=0, column=0, sticky="nsew")
    
        # Crear la Scrollbar
        scrollbar = tk.Scrollbar(revisar_frame, orient="vertical", command=tree.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        tree.configure(yscrollcommand=scrollbar.set)
    
        # Ordenar expedientes alfabéticamente antes de añadirlos
        expedientes_ordenados = sorted(planes_trabajo.keys())
        for expediente in expedientes_ordenados:
            tree.insert("", "end", values=(expediente,))

        
    def informe_de_trabajo(self):
        self.clear_body_frame()  # Limpiar el área principal antes de mostrar el formulario
         # Título del formulario
        tk.Label(
            self.body_frame,
            text="Crear Nuevo Informe de Seguimiento",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#800000"
        ).pack(pady=10)

        frame_exp = tk.Frame(self.body_frame, bg="white", width=300, height=200, highlightbackground= "black", highlightthickness=2)
        frame_exp.pack(padx=10, pady=10)
        canvas = tk.Canvas(frame_exp, bg="white")
        scrollbar = tk.Scrollbar(frame_exp, orient="vertical", command=canvas.yview)
        contenido_frame = tk.Frame(canvas, bg="white")

        contenido_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=contenido_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Posicionar Canvas y Scrollbar dentro del Frame
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # -------------------- CARGAR Y FILTRAR DATOS --------------------
        planes = self.cargar_planes()
        planes_filtrados = [p for p in planes if str(p.get("Número de Expediente", "")) == str(numero_expediente)]

        # Mostrar los datos en etiquetas dentro del Frame
        if planes_filtrados:
            for plan in planes_filtrados:
                texto = f"Expediente: {plan.get('Número de Expediente', 'N/A')}\n"
                texto += f"Descripción: {plan.get('Descripción', 'Sin descripción')}\n"
                tk.Label(contenido_frame, text=texto, bg="white", justify="left", anchor="w").pack(fill="x", padx=5, pady=5)
        else:
            tk.Label(contenido_frame, text="No se encontraron planes de trabajo.", bg="white").pack(pady=10)
    

    def crear_plan_trabajo(self):
        
    # Verificar si el archivo 'configuracion.json' existe
        if not os.path.exists("configuracion.json"):
            messagebox.showerror(
                "Configuración faltante",
                "El archivo 'configuracion.json' no existe. Por favor, llene la configuración antes de continuar."
            )
            return  # No continuar con la ejecución de la función
        """Muestra el formulario completo en el área principal (body_frame) para crear un nuevo plan de trabajo."""
        self.clear_body_frame()  # Limpiar el área principal antes de mostrar el formulario

        # Título del formulario
        tk.Label(
            self.body_frame,
            text="Crear Nuevo Plan de Trabajo",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#800000"
        ).pack(pady=10)

        # Contenedor para organizar los campos del formulario
        form_frame = tk.Frame(self.body_frame, bg="#FFFFFF")
        form_frame.pack(pady=10, padx=10, fill="both", expand=True)

        # Campos del formulario
        tk.Label(form_frame, text="Fecha (dd-mm-yyyy):", bg="#FFFFFF").grid(row=0, column=0, sticky="w", pady=5)
        fecha = tk.Entry(form_frame)
        fecha.grid(row=0, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Unidad:", bg="#FFFFFF").grid(row=1, column=0, sticky="w", pady=5)
        unidad = tk.Entry(form_frame)
        unidad.grid(row=1, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Número de Expediente:", bg="#FFFFFF").grid(row=2, column=0, sticky="w", pady=5)
        expediente = tk.Entry(form_frame)
        expediente.grid(row=2, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Nombre del paciente:", bg="#FFFFFF").grid(row=3, column=0, sticky="w", pady=5)
        nombre = tk.Entry(form_frame)
        nombre.grid(row=3, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Edad:", bg="#FFFFFF").grid(row=4, column=0, sticky="w", pady=5)
        edad = tk.Entry(form_frame)
        edad.grid(row=4, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Sexo (M/F):", bg="#FFFFFF").grid(row=5, column=0, sticky="w", pady=5)
        sexo = tk.Entry(form_frame)
        sexo.grid(row=5, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Fecha de Nacimiento (dd-mm-yyyy):", bg="#FFFFFF").grid(row=6, column=0, sticky="w", pady=5)
        nacimiento = tk.Entry(form_frame)
        nacimiento.grid(row=6, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Diagnóstico:", bg="#FFFFFF").grid(row=7, column=0, sticky="w", pady=5)
        diagnostico = tk.Entry(form_frame)
        diagnostico.grid(row=7, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Área de Intervención:", bg="#FFFFFF").grid(row=8, column=0, sticky="w", pady=5)
        area_intervencion = tk.Entry(form_frame)
        area_intervencion.grid(row=8, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Número de sesiones por semana:", bg="#FFFFFF").grid(row=9, column=0, sticky="w", pady=5)
        sesiones = tk.Entry(form_frame)
        sesiones.grid(row=9, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Reforzadores Comestibles:", bg="#FFFFFF").grid(row=10, column=0, sticky="w", pady=5)
        reforzadores_comestibles = tk.Entry(form_frame)
        reforzadores_comestibles.grid(row=10, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Reforzadores Tangibles:", bg="#FFFFFF").grid(row=11, column=0, sticky="w", pady=5)
        reforzadores_tangibles = tk.Entry(form_frame)
        reforzadores_tangibles.grid(row=11, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Reforzadores Sociales:", bg="#FFFFFF").grid(row=12, column=0, sticky="w", pady=5)
        reforzadores_sociales = tk.Entry(form_frame)
        reforzadores_sociales.grid(row=12, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Observaciones clínicas:", bg="#FFFFFF").grid(row=13, column=0, sticky="w", pady=5)
        observaciones = tk.Entry(form_frame)
        observaciones.grid(row=13, column=1, pady=5, padx=10)

        tk.Label(form_frame, text="Programas seleccionados (1,2,5,10...):", bg="#FFFFFF").grid(row=14, column=0, sticky="w", pady=5)
        programas_seleccionados = tk.Entry(form_frame)
        programas_seleccionados.grid(row=14, column=1, pady=5, padx=10)
        

        # Botón de Directorio para abrir la función F3
        boton_directorio = tk.Button(
            form_frame,
            text="Directorio",
            command=lambda: seleccionar_programas_plan(programas_seleccionados),
            bg="#A52A2A",
            fg="white",
            relief="flat",
            activebackground="#800000",
            activeforeground="white"
        )
        boton_directorio.grid(row=14, column=2, pady=5, padx=10)
        def guardar_plan():
            datos = {
                "Fecha": fecha.get(),
                "Unidad": unidad.get(),
                "Número de Expediente": expediente.get(),
                "Nombre del Paciente": nombre.get(),
                "Edad": edad.get(),
                "Sexo": sexo.get(),
                "Fecha de Nacimiento": nacimiento.get(),
                "Diagnóstico": diagnostico.get(),
                "Área de Intervención": area_intervencion.get(),
                "Sesiones por Semana": sesiones.get(),
                "Reforzadores Comestibles": reforzadores_comestibles.get(),
                "Reforzadores Tangibles": reforzadores_tangibles.get(),
                "Reforzadores Sociales": reforzadores_sociales.get(),
                "Observaciones Clínicas": observaciones.get(),
                "Programas Seleccionados": [int(p) for p in programas_seleccionados.get().split(",")]
            }
            guardar_informacion(expediente.get(), datos)
            crear_documento(datos)
            messagebox.showinfo("Guardado", "El plan de trabajo ha sido guardado correctamente.")
            self.clear_body_frame()     
        # Botón de Guardar
        boton_guardar = tk.Button(
            form_frame,
            text="Guardar plan",
            command= guardar_plan,  # Ejecutar guardar_plan cuando se presione el botón
            bg="#A52A2A",
            fg="white",
            relief="flat",
            activebackground="#800000",
            activeforeground="white"
        )
        boton_guardar.grid(row=16, column=2, pady=5, padx=10)     

        # Configurar las columnas del formulario para que se ajusten
        form_frame.columnconfigure(1, weight=1)


        
    def revisar_planes_trabajo(self):
        archivo_planes = "planes_trabajo.json"  # Usamos el archivo centralizado
    
        if not os.path.exists(archivo_planes):
            messagebox.showinfo("Error", "No se han encontrado planes de trabajo guardados.")
            return
    
        with open(archivo_planes, "r") as f:
            planes_trabajo = json.load(f)
    
        if not planes_trabajo:
            messagebox.showinfo("Error", "No hay planes de trabajo disponibles.")
            return
    
        self.clear_body_frame()
    
        # Título
        tk.Label(
            self.body_frame,
            text="Revisar planes de trabajo",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#800000"
        ).grid(row=0, column=0, pady=10)
    
        # Configurar el grid para el body_frame
        self.body_frame.grid_rowconfigure(1, weight=1)
        self.body_frame.grid_columnconfigure(0, weight=1)
    
        # Contenedor principal
        revisar_frame = tk.Frame(self.body_frame, bg="#FFFFFF")
        revisar_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
    
        # Configurar el grid para revisar_frame
        revisar_frame.grid_rowconfigure(0, weight=1)
        revisar_frame.grid_columnconfigure(0, weight=1)
    
        # Crear el Treeview
        tree = ttk.Treeview(revisar_frame, columns=("Expediente"), show="headings", height=10)
        tree.heading("Expediente", text="Expediente")
        tree.column("Expediente", anchor="w")
        tree.grid(row=0, column=0, sticky="nsew")
    
        # Crear la Scrollbar
        scrollbar = tk.Scrollbar(revisar_frame, orient="vertical", command=tree.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        tree.configure(yscrollcommand=scrollbar.set)
    
        # Ordenar expedientes alfabéticamente antes de añadirlos
        expedientes_ordenados = sorted(planes_trabajo.keys())
        for expediente in expedientes_ordenados:
            tree.insert("", "end", values=(expediente,))
            
            # Función para seleccionar programas
            def seleccionar_programas():
                # Crear una nueva ventana para seleccionar programas
                ventana_seleccion = tk.Toplevel()
                ventana_seleccion.title("Seleccionar Programas")
                ventana_seleccion.geometry("400x300")  # Tamaño de la ventana para que sea más visible
            
                # Frame contenedor con scrollbar
                frame_contenedor = tk.Frame(ventana_seleccion)
                frame_contenedor.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
            
                # Configuración del grid de la ventana
                ventana_seleccion.grid_rowconfigure(0, weight=1)
                ventana_seleccion.grid_columnconfigure(0, weight=1)
            
                # Canvas y scrollbar
                canvas = tk.Canvas(frame_contenedor)
                scrollbar = tk.Scrollbar(frame_contenedor, orient="vertical", command=canvas.yview)
                frame_scrollable = tk.Frame(canvas)
            
                # Configuración de scrollbar
                frame_scrollable.bind(
                    "<Configure>",
                    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
                )
            
                # Colocar el frame scrollable en el canvas
                canvas.create_window((0, 0), window=frame_scrollable, anchor="nw")
                canvas.configure(yscrollcommand=scrollbar.set)
            
                # Colocar canvas y scrollbar en el grid
                canvas.grid(row=0, column=0, sticky="nsew")
                scrollbar.grid(row=0, column=1, sticky="ns")
            
                # Expandir el frame contenedor y el canvas para ajustarse al tamaño de la ventana
                frame_contenedor.grid_rowconfigure(0, weight=1)
                frame_contenedor.grid_columnconfigure(0, weight=1)
            
                seleccionados = []  # Lista para almacenar los programas seleccionados
            
                for i, (programa_id, programa_info) in enumerate(programas.items()):
                    var = tk.IntVar()  # Variable que se usará para marcar el checkbox
                    check = tk.Checkbutton(frame_scrollable, text=programa_info["Nombre"], variable=var)
                    check.grid(row=i, column=0, sticky='w', pady=2)  # Usar grid para colocar los checkboxes
                    seleccionados.append((programa_id, var))  # Guardar el id y el estado de selección
            
                # Variable para almacenar los programas seleccionados
                seleccion_final = []
            
                # Función que se ejecuta cuando se confirman los programas seleccionados
                def confirmar_seleccion():
                    global seleccion_final
                    # Convertir los IDs seleccionados a enteros
                    seleccion_final = [int(id) for id, var in seleccionados if var.get() == 1]
            
                    # Actualizar el campo "Programas Seleccionados" en la ventana de detalles
                    if "Programas Seleccionados" in campos_modificados:
                        # Limpiar el campo antes de insertar los nuevos valores
                        campos_modificados["Programas Seleccionados"].delete(0, tk.END)
                        # Insertar los IDs de los programas seleccionados separados por comas (convertidos a strings para mostrar)
                        campos_modificados["Programas Seleccionados"].insert(0, ','.join(map(str, seleccion_final)))
            
                    # Cerrar la ventana de selección después de confirmar
                    ventana_seleccion.destroy()
            
                # Botón para confirmar la selección
                boton_confirmar = tk.Button(ventana_seleccion, text="Confirmar", command=confirmar_seleccion)
                boton_confirmar.grid(row=1, column=0, pady=10)  # Usar grid para colocar el botón
            
                ventana_seleccion.mainloop()
            
        
    
            # Función para mostrar detalles
            def mostrar_detalles():
                seleccion = tree.selection()
                if not seleccion:
                    messagebox.showinfo("Error", "Por favor, seleccione un plan de trabajo.")
                    return
            
                expediente_seleccionado = tree.item(seleccion, "values")[0]
                datos_plan = planes_trabajo[expediente_seleccionado]
            
                # Crear un diccionario para almacenar los campos modificados
                campos_modificados = {}   
            
                self.clear_body_frame()             
            
                # Título
                tk.Label(
                    self.body_frame,
                    text="Detalles del plan seleccionado",
                    font=("Arial", 16, "bold"),
                    bg="#FFFFFF",
                    fg="#000000"
                ).grid(row=0, column=0, columnspan=2, pady=10)
            
                # Contenedor para los detalles
                detalles_frame = tk.Frame(self.body_frame, bg="#FFFFFF")
                detalles_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
            
                self.body_frame.grid_rowconfigure(1, weight=1)
                self.body_frame.grid_columnconfigure(0, weight=1)
            
                # Canvas con barra de desplazamiento
                canvas = tk.Canvas(detalles_frame, bg="#FFFFFF")
                scrollbar = tk.Scrollbar(detalles_frame, orient="vertical", command=canvas.yview)
                scrollable_frame = tk.Frame(canvas, bg="#FFFFFF")
                scrollable_frame.grid_columnconfigure(0, weight=1)  # Para las etiquetas
                scrollable_frame.grid_columnconfigure(1, weight=2)  # Para los campos de entrada
            
                scrollable_frame.bind(
                    "<Configure>",
                    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
                )
            
                canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
                canvas.configure(yscrollcommand=scrollbar.set)
            
                canvas.grid(row=0, column=0, sticky="nsew")
                scrollbar.grid(row=0, column=1, sticky="ns")
            
                detalles_frame.grid_rowconfigure(0, weight=1)
                detalles_frame.grid_columnconfigure(0, weight=1)
            
                # Crear una fila para cada campo
                for idx, (key, value) in enumerate(datos_plan.items()):
                    # Etiqueta para el nombre del campo
                    etiqueta = tk.Label(
                        scrollable_frame, 
                        text=key, 
                        bg="#FFFFFF",
                        anchor="e",
                        justify="right",
                        width=20
                    )
                    etiqueta.grid(row=idx, column=0, padx=10, pady=5, sticky="e")
            
                    # Campo de entrada para el valor
                    entry = tk.Entry(scrollable_frame, width=40)
                    entry.insert(0, value if not isinstance(value, list) else ','.join(map(str, value)))
                    entry.grid(row=idx, column=1, padx=10, pady=5, sticky="w")
            
                    # Almacenar los campos modificados
                    campos_modificados[key] = entry
                    
                # Mostrar los programas seleccionados en el plan de trabajo
                row_offset = idx + 1  # Iniciar desde la fila siguiente al último índice utilizado
                etiqueta_programas = tk.Label(
                    scrollable_frame, 
                    text="Programas Seleccionados:", 
                    font=("Arial", 10, "bold"),
                    bg="#FFFFFF"
                )
                etiqueta_programas.grid(row=row_offset, column=0, columnspan=2, padx=10, pady=10, sticky="n")
                
                # Iterar sobre los programas seleccionados
                for i, programa in enumerate(datos_plan.get("Programas Seleccionados", []), start=row_offset + 1):
                    programa_int = int(programa)  # Convertir a entero si es necesario
                    nombre_programa = programas.get(programa_int, {}).get("Nombre", f"Programa {programa} no encontrado")
                    
                    # Etiqueta para cada programa centrada
                    etiqueta_programa = tk.Label(
                        scrollable_frame, 
                        text=f"- {nombre_programa}", 
                        font=("Arial", 9),
                        bg="white", 
                        anchor="center",
                        justify="center"
                    )
                    etiqueta_programa.grid(row=i, column=0, columnspan=2, padx=20, pady=5, sticky="n")
                    
            
                    # Función para guardar los cambios
                    def guardar_todos_los_cambios():
                        # Iterar sobre los campos modificados y guardar sus nuevos valores
                        for key, entry in campos_modificados.items():
                            nuevo_valor = entry.get()
                            if key == "Programas Seleccionados":  # Si es la lista de programas, procesarlo como una lista
                                datos_plan[key] = nuevo_valor.split(',')  # Convertir la cadena en lista
                            else:
                                datos_plan[key] = nuevo_valor  # Guardar el nuevo valor del campo
                    
                        # Guardar los datos actualizados en el diccionario de planes de trabajo
                        planes_trabajo[expediente_seleccionado] = datos_plan
                    
                        # Guardar los cambios a nivel de almacenamiento
                        guardar_planes_trabajo(planes_trabajo)
                    
                        messagebox.showinfo("Éxito", "Los cambios se han guardado correctamente.")
                        self.clear_body_frame()  # Limpiar el área principal antes de mostrar el formulario
                        
                    def modificar_programas():
                        seleccion_actual = datos_plan.get("Programas Seleccionados", [])
                        seleccion_actual = set(map(int, seleccion_actual))  # Convertir a conjunto de enteros
                    
                        # Crear una nueva ventana para seleccionar programas
                        ventana_seleccion = tk.Toplevel()
                        ventana_seleccion.title("Seleccionar Programas")
                        ventana_seleccion.geometry("400x300")
                    
                        # Frame contenedor con scrollbar
                        frame_contenedor = tk.Frame(ventana_seleccion, bg="white")
                        frame_contenedor.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
                    
                        ventana_seleccion.grid_rowconfigure(0, weight=1)
                        ventana_seleccion.grid_columnconfigure(0, weight=1)
                    
                        # Canvas y scrollbar
                        canvas = tk.Canvas(frame_contenedor, bg="white")
                        scrollbar = tk.Scrollbar(frame_contenedor, orient="vertical", command=canvas.yview)
                        frame_scrollable = tk.Frame(canvas, bg="white")
                    
                        frame_scrollable.bind(
                            "<Configure>",
                            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
                        )
                    
                        canvas.create_window((0, 0), window=frame_scrollable, anchor="nw")
                        canvas.configure(yscrollcommand=scrollbar.set)
                    
                        canvas.grid(row=0, column=0, sticky="nsew")
                        scrollbar.grid(row=0, column=1, sticky="ns")
                    
                        frame_contenedor.grid_rowconfigure(0, weight=1)
                        frame_contenedor.grid_columnconfigure(0, weight=1)
                    
                        # Lista para almacenar las opciones seleccionadas
                        seleccionados = []
                    
                        # Crear el checklist
                        for programa_id, programa_data in programas.items():
                            var = tk.IntVar(value=1 if programa_id in seleccion_actual else 0)  # Marcar si está seleccionado
                            chk = tk.Checkbutton(
                                frame_scrollable,
                                text=programa_data["Nombre"],
                                variable=var,
                                bg="#FFFFFF",
                                anchor="w",
                                justify="left"
                            )
                            chk.pack(fill="x", padx=5, pady=2, anchor="w")
                            seleccionados.append((programa_id, var))
                    
                        # Función para confirmar la selección
                        def confirmar_seleccion():
                            seleccion_final = [
                                str(programa_id) for programa_id, var in seleccionados if var.get() == 1
                            ]
                    
                            # Actualizar el campo "Programas Seleccionados" en `campos_modificados`
                            if "Programas Seleccionados" in campos_modificados:
                                entry = campos_modificados["Programas Seleccionados"]
                                entry.delete(0, tk.END)  # Limpiar el campo de entrada
                                entry.insert(0, ','.join(seleccion_final))  # Insertar los programas seleccionados
                    
                            # Actualizar los datos en el plan actual
                            datos_plan["Programas Seleccionados"] = seleccion_final
                    
                            messagebox.showinfo(
                                "Selección Confirmada", f"Programas seleccionados: {', '.join(seleccion_final)}"
                            )
                            ventana_seleccion.destroy()
                    
                        # Función para borrar toda la selección
                        def borrar_seleccion():
                            for _, var in seleccionados:
                                var.set(0)  # Desmarcar todas las opciones
                    
                        # Botón para confirmar la selección
                        boton_confirmar = tk.Button(
                            ventana_seleccion,
                            text="Confirmar Selección",
                            command=confirmar_seleccion,
                            bg="#A52A2A",
                            fg="white",
                            font=("Arial", 12)
                        )
                        boton_confirmar.grid(row=1, column=0, pady=10, sticky="ew")
                    
                        # Botón para borrar la selección
                        boton_borrar = tk.Button(
                            ventana_seleccion,
                            text="Borrar Selección",
                            command=borrar_seleccion,
                            bg="#007BFF",
                            fg="white",
                            font=("Arial", 12)
                        )
                        boton_borrar.grid(row=2, column=0, pady=10, sticky="ew")
                    
                        ventana_seleccion.mainloop()
                    
                 
                # Contenedor para los botones
                botones_frame = tk.Frame(self.body_frame, bg="#FFFFFF")
                botones_frame.grid(row=row_offset + 2, column=0, columnspan=2, pady=20, sticky="ew")
            
                # Configuración para distribuir botones uniformemente
                botones_frame.grid_columnconfigure(0, weight=1)
                botones_frame.grid_columnconfigure(1, weight=1)
                botones_frame.grid_columnconfigure(2, weight=1)
            
                # Botón para guardar todos los cambios
                boton_guardar = tk.Button(
                    botones_frame, 
                    text="Guardar Cambios", 
                    command=guardar_todos_los_cambios, 
                    width=20
                )
                boton_guardar.grid(row=0, column=0, padx=10, pady=5)
            
                # Botón para modificar los programas
                boton_modificar_programas = tk.Button(
                    botones_frame, 
                    text="Modificar Programas", 
                    command=modificar_programas, 
                    width=20
                )
                boton_modificar_programas.grid(row=0, column=1, padx=10, pady=5)
            
                # Botón para crear un documento
                boton_crear_documento = tk.Button(
                    botones_frame, 
                    text="Crear Documento", 
                    command=lambda: usar_crear_documento(expediente_seleccionado), 
                    width=20
                )
                boton_crear_documento.grid(row=0, column=2, padx=10, pady=5)

            # Botón para ver detalles
            boton_detalles = tk.Button(
                self.body_frame,
                text="Ver Detalles",
                command=mostrar_detalles,
                bg="#A52A2A",
                fg="white",
                font=("Arial", 12),
                relief="flat",
                activebackground="#800000",
                activeforeground="white"
            )
            boton_detalles.grid(row=2, column=0, pady=10, sticky="ew")



    def abrir_directorio_programas(self):
        """Muestra el directorio de programas en el área principal."""
        self.clear_body_frame()
        
        # Encabezado estilizado
        header_label = tk.Label(
            self.body_frame,
            text="Directorio",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#800000"
        )
        header_label.pack(pady=10)
    
        # Marco para contener el Treeview y la barra de desplazamiento
        frame_tree = tk.Frame(self.body_frame, bg="#FFFFFF")
        frame_tree.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Barra de desplazamiento vertical
        scrollbar = ttk.Scrollbar(frame_tree, orient="vertical")
        scrollbar.pack(side="right", fill="y")
        
        # Crear Treeview para mostrar programas
        self.tree = ttk.Treeview(
            frame_tree,
            columns=("ID", "Nombre", "Objetivo"),
            show="headings",
            selectmode="browse",
            yscrollcommand=scrollbar.set,
        )
        self.tree.heading("ID", text="ID Programa")
        self.tree.heading("Nombre", text="Nombre")
        self.tree.heading("Objetivo", text="Objetivo (Resumen)")
        
        self.tree.column("ID", width=80, anchor="center")
        self.tree.column("Nombre", width=200, anchor="w")
        self.tree.column("Objetivo", width=600, anchor="w")
        self.tree.pack(fill="both", expand=True)
        
        scrollbar.config(command=self.tree.yview)
        
        # Llenar Treeview con datos
        self.actualizar_lista_programas()
    
        # Vincular tecla Enter para mostrar detalles
        self.tree.bind("<Return>", lambda event: self.mostrar_detalles(self.tree))
        self.tree.bind("<Double-1>", lambda event: self.mostrar_detalles(self.tree))  # Abrir con doble clic también
        
        # Botón para abrir detalles
        boton_abrir = tk.Button(
            self.body_frame,
            text="Abrir detalles",
            command=lambda: self.mostrar_detalles(self.tree),
            font=("Arial", 12),
            bg="#8B0000",
            fg="white",
            relief="flat",
            activebackground="#A52A2A",
            activeforeground="white",
        )
        boton_abrir.pack(pady=5)
    
        # Botón para agregar nuevos programas
        boton_agregar = tk.Button(
            self.body_frame,
            text="Agregar nuevo programa",
            command=self.agregar_programa,
            font=("Arial", 12),
            bg="#8B0000",
            fg="white",
            relief="flat",
            activebackground="#A52A2A",
            activeforeground="white",
        )
        boton_agregar.pack(pady=5)
    
        # Botón para borrar programas
        boton_borrar = tk.Button(
            self.body_frame,
            text="Borrar programa seleccionado",
            command=self.borrar_programa,
            font=("Arial", 12),
            bg="#8B0000",
            fg="white",
            relief="flat",
            activebackground="#A52A2A",
            activeforeground="white",
        )
        boton_borrar.pack(pady=5)
    
    def actualizar_lista_programas(self):
        """Actualiza la lista de programas en el Treeview."""
        for item in self.tree.get_children():
            self.tree.delete(item)
        for programa_id, detalles in programas.items():
            nombre_programa = detalles.get("Nombre", "Sin nombre")
            objetivo_resumen = detalles.get("Objetivo", "Sin objetivo")[:50] + "..."
            self.tree.insert("", tk.END, values=(programa_id, nombre_programa, objetivo_resumen))
    
    def agregar_programa(self):
        """Abre una ventana para agregar un nuevo programa."""
        agregar_window = tk.Toplevel(self.root)
        agregar_window.title("Agregar Nuevo Programa")
        agregar_window.geometry("600x600")
        agregar_window.configure(bg="#FFFFFF")

        # Títulos
        tk.Label(
            agregar_window,
            text="Agregar Nuevo Programa",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#000000"
        ).pack(pady=10)

        # Campo: Nombre
        tk.Label(
            agregar_window,
            text="Nombre:",
            font=("Arial", 14),
            bg="#FFFFFF"
        ).pack(anchor="w", padx=10, pady=(10, 0))
        entry_nombre = tk.Entry(agregar_window, font=("Arial", 12))
        entry_nombre.pack(fill="x", padx=10, pady=5)

        # Campo: Objetivo
        tk.Label(
            agregar_window,
            text="Objetivo:",
            font=("Arial", 14),
            bg="#FFFFFF"
        ).pack(anchor="w", padx=10, pady=(10, 0))
        text_objetivo = tk.Text(agregar_window, font=("Arial", 12), height=4)
        text_objetivo.pack(fill="x", padx=10, pady=5)

        # Campo: Procedimiento
        tk.Label(
            agregar_window,
            text="Procedimiento:",
            font=("Arial", 14),
            bg="#FFFFFF"
        ).pack(anchor="w", padx=10, pady=(10, 0))
        text_procedimiento = tk.Text(agregar_window, font=("Arial", 12), height=6)
        text_procedimiento.pack(fill="x", padx=10, pady=5)

        # Campo: Ayudas
        tk.Label(
            agregar_window,
            text="Ayudas:",
            font=("Arial", 14),
            bg="#FFFFFF"
        ).pack(anchor="w", padx=10, pady=(10, 0))
        text_ayudas = tk.Text(agregar_window, font=("Arial", 12), height=4)
        text_ayudas.pack(fill="x", padx=10, pady=5)

        def guardar_nuevo_programa():
            nuevo_id = len(programas) + 1  # Generar nuevo ID
            programas[nuevo_id] = {
                "Nombre": entry_nombre.get(),
                "Objetivo": text_objetivo.get("1.0", "end").strip(),
                "Procedimiento": text_procedimiento.get("1.0", "end").strip(),
                "Ayudas": text_ayudas.get("1.0", "end").strip(),
            }

            # Guardar en el archivo JSON
            with open("programas.json", "w") as file:
                json.dump(programas, file, indent=4)

            agregar_window.destroy()
            tk.messagebox.showinfo("Éxito", f"El programa {nuevo_id} ha sido agregado.")
                # Llamar a la función para actualizar la lista
            self.actualizar_lista_programas()
            # Botón para guardar el programa
        tk.Button(
            agregar_window,
            text="Guardar Programa",
            command=guardar_nuevo_programa,
            bg="#A52A2A",
            fg="white",
            font=("Arial", 12),
        ).pack(pady=10)

        # Botón para cancelar
        tk.Button(
            agregar_window,
            text="Cancelar",
            command=agregar_window.destroy,
            bg="#A52A2A",
            fg="white",
            font=("Arial", 12),
        ).pack(pady=10)
    
    def borrar_programa(self):
        """Borra el programa seleccionado en el Treeview."""
        selected_item = self.tree.selection()
        if not selected_item:
            tk.messagebox.showwarning("Advertencia", "Seleccione un programa para borrar.")
            return
    
        programa_id = int(self.tree.item(selected_item, "values")[0])
        confirmar = tk.messagebox.askyesno("Confirmar", f"¿Está seguro de que desea borrar el programa {programa_id}?")
        if confirmar:
            programas.pop(programa_id, None)
            self.actualizar_lista_programas()
            tk.messagebox.showinfo("Éxito", f"El programa {programa_id} ha sido borrado.")
        
    def mostrar_detalles(self, tree):
        """Muestra y permite modificar los detalles del programa seleccionado."""
        item = tree.selection()
        if item:
            programa_id = tree.item(item)["values"][0]
            detalles = programas.get(programa_id, {})
            
            # Obtener los datos del programa
            nombre = detalles.get("Nombre", "Sin nombre")
            objetivo = detalles.get("Objetivo", "Sin objetivo")
            procedimiento = detalles.get("Procedimiento", "Sin procedimiento definido")
            ayudas = detalles.get("Ayudas", "Sin ayudas disponibles")
    
            # Crear ventana emergente para mostrar detalles
            detalles_window = tk.Toplevel(self.root)
            detalles_window.title(f"Detalles del Programa {programa_id}")
            detalles_window.geometry("600x600")
            detalles_window.configure(bg="#FFFFFF")
    
            # Encabezado: Nombre del programa
            tk.Label(
                detalles_window,
                text=nombre,
                font=("Arial", 16, "bold"),
                bg="#FFFFFF",
                fg="#000000"
            ).pack(pady=10)
    
            # Sección: Objetivo
            tk.Label(
                detalles_window,
                text="Objetivo:",
                font=("Arial", 14, "bold"),
                bg="#FFFFFF",
                fg="#800000"
            ).pack(anchor="w", padx=10, pady=(10, 0))
            text_objetivo = tk.Text(
                detalles_window,
                wrap="word",
                font=("Arial", 12),
                bg="#F5F5F5",
                fg="#000000",
                height=4
            )
            text_objetivo.insert("1.0", objetivo)
            text_objetivo.pack(fill="x", padx=10, pady=5)
    
            # Sección: Procedimiento
            tk.Label(
                detalles_window,
                text="Procedimiento:",
                font=("Arial", 14, "bold"),
                bg="#FFFFFF",
                fg="#800000"
            ).pack(anchor="w", padx=10, pady=(10, 0))
            text_procedimiento = tk.Text(
                detalles_window,
                wrap="word",
                font=("Arial", 12),
                bg="#F5F5F5",
                fg="#000000",
                height=6
            )
            text_procedimiento.insert("1.0", procedimiento)
            text_procedimiento.pack(fill="x", padx=10, pady=5)
    
            # Sección: Ayudas
            tk.Label(
                detalles_window,
                text="Ayudas:",
                font=("Arial", 14, "bold"),
                bg="#FFFFFF",
                fg="#800000"
            ).pack(anchor="w", padx=10, pady=(10, 0))
            text_ayudas = tk.Text(
                detalles_window,
                wrap="word",
                font=("Arial", 12),
                bg="#F5F5F5",
                fg="#000000",
                height=4
            )
            text_ayudas.insert("1.0", ayudas)
            text_ayudas.pack(fill="x", padx=10, pady=5)
    
            # Función para guardar los cambios en el archivo JSON
            def guardar_cambios():
                # Actualizar el diccionario con los cambios
                programas[programa_id]["Objetivo"] = text_objetivo.get("1.0", "end").strip()
                programas[programa_id]["Procedimiento"] = text_procedimiento.get("1.0", "end").strip()
                programas[programa_id]["Ayudas"] = text_ayudas.get("1.0", "end").strip()
            
                # Guardar los cambios en el archivo programas.json
                with open("programas.json", "w") as file:
                    json.dump(programas, file, indent=4)
            
                # Cerrar la ventana de detalles
                detalles_window.destroy()
                
                # Mostrar mensaje de confirmación
                messagebox.showinfo("Guardado", f"Los cambios en el programa {programa_id} se han guardado.")
            
            # Botón para guardar cambios
            tk.Button(
                detalles_window,
                text="Guardar Cambios",
                command=guardar_cambios,
                bg="#800000",
                fg="white",
                font=("Arial", 12),
            ).pack(pady=10)
            
            # Botón para cerrar la ventana
            tk.Button(
                detalles_window,
                text="Cerrar",
                command=detalles_window.destroy,
                bg="#800000",
                fg="white",
                font=("Arial", 12),
            ).pack(pady=10)
    
    # Función para borrar un programa y actualizar el archivo JSON
    def borrar_programa(self):
        """Borra el programa seleccionado en el Treeview y actualiza el archivo JSON."""
        selected_item = self.tree.selection()
        if not selected_item:
            tk.messagebox.showwarning("Advertencia", "Seleccione un programa para borrar.")
            return
    
        programa_id = int(self.tree.item(selected_item, "values")[0])
        confirmar = tk.messagebox.askyesno("Confirmar", f"¿Está seguro de que desea borrar el programa {programa_id}?")
        if confirmar:
            # Eliminar el programa del diccionario
            programas.pop(programa_id, None)
    
            # Guardar los cambios en el archivo JSON después de borrar
            with open("programas.json", "w") as file:
                json.dump(programas, file, indent=4)
    
            # Actualizar la lista de programas en la interfaz
            self.actualizar_lista_programas()
    
            # Mostrar mensaje de confirmación
            tk.messagebox.showinfo("Éxito", f"El programa {programa_id} ha sido borrado.")

            
            
    def abrir_configuracion(self):
        """Muestra el formulario de configuración en el área principal."""
        self.clear_body_frame()  # Limpiar el área principal
        configuracion_actual = cargar_configuracion()
    
        # Encabezado estilizado
        header_label = tk.Label(
            self.body_frame,
            text="Configuración del Sistema",
            font=("Arial", 16, "bold"),
            bg="#FFFFFF",
            fg="#800000"
        )
        header_label.pack(pady=10)
    
        # Marco para contener los campos
        form_frame = tk.Frame(self.body_frame, bg="#F5F5F5", padx=20, pady=20)
        form_frame.pack(pady=20, padx=10, fill="both", expand=True)
    
        # Campo: Unidad
        tk.Label(form_frame, text="Unidad:", font=("Arial", 12), bg="#F5F5F5").grid(row=0, column=0, sticky="w", pady=5, padx=5)
        entrada_unidad = tk.Entry(form_frame, font=("Arial", 12), width=30)
        entrada_unidad.grid(row=0, column=1, pady=5, padx=5)
        entrada_unidad.insert(0, configuracion_actual.get("unidad", ""))
    
        # Campo: Nombre del Terapeuta
        tk.Label(form_frame, text="Nombre del Terapeuta:", font=("Arial", 12), bg="#F5F5F5").grid(row=1, column=0, sticky="w", pady=5, padx=5)
        entrada_nombre_terapeuta = tk.Entry(form_frame, font=("Arial", 12), width=30)
        entrada_nombre_terapeuta.grid(row=1, column=1, pady=5, padx=5)
        entrada_nombre_terapeuta.insert(0, configuracion_actual.get("nombre_terapeuta", ""))
    
        # Campo: Cédula Profesional
        tk.Label(form_frame, text="Cédula Profesional:", font=("Arial", 12), bg="#F5F5F5").grid(row=2, column=0, sticky="w", pady=5, padx=5)
        entrada_cedula_profesional = tk.Entry(form_frame, font=("Arial", 12), width=30)
        entrada_cedula_profesional.grid(row=2, column=1, pady=5, padx=5)
        entrada_cedula_profesional.insert(0, configuracion_actual.get("cedula_profesional", ""))
    
        # Botón para guardar la configuración
        boton_guardar = tk.Button(
            self.body_frame,
            text="Guardar Configuración",
            font=("Arial", 12),
            bg="#A52A2A",
            fg="white",
            command=lambda: guardar_configuracion(
                entrada_unidad.get(),
                entrada_nombre_terapeuta.get(),
                entrada_cedula_profesional.get()
            )
        )
        boton_guardar.pack(pady=20)
    
        # Agregar marco de separación visual
        separator = tk.Frame(self.body_frame, height=2, bd=1, relief="sunken", bg="#D3D3D3")
        separator.pack(fill="x", pady=10)
    

    def run(self):
        self.root.mainloop()

# Ejecutar el menú principal
if __name__ == "__main__":
    menu = ModernMainMenu()
    menu.run()
