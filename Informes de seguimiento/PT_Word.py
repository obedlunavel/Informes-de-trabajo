import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.shared import Pt, RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import tkinter.simpledialog as simpledialog
import sys

# Diccionario con los programas de intervención
programas = {}

# Verificar si el archivo 'programas.json' existe
if os.path.exists("programas.json"):
    # Si el archivo existe, cargar su contenido en programas
    with open("programas.json", "r") as file:
        programas_json = json.load(file)
    
    # Limpiar y actualizar el diccionario `programas` con los datos del archivo
    programas.clear()
    programas.update({int(k): v for k, v in programas_json.items()})
else:
    # Si el archivo no existe, crearlo con un diccionario vacío
    with open("programas.json", "w") as file:
        json.dump(programas, file, indent=4)
class GeneradorDocumento:
    def __init__(self, config_path='configuracion.json'):
        """
        Inicializa la clase GeneradorDocumento.

        Parámetros:
            config_path (str): Ruta al archivo de configuración JSON.
        """
        self.config_path = config_path
        self.configuracion = self._cargar_configuracion()

    @staticmethod
    def resource_path(relative_path):
        """
        Obtiene la ruta absoluta al recurso, funciona para desarrollo y para PyInstaller.

        Parámetros:
            relative_path (str): Ruta relativa al recurso.

        Retorna:
            str: Ruta absoluta al recurso.
        """
        try:
            # PyInstaller crea una carpeta temporal y almacena la ruta en _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)

    def _cargar_configuracion(self):
        """
        Carga la configuración desde el archivo JSON.

        Retorna:
            dict: Diccionario con la configuración.
        """
        with open(self.config_path, 'r') as file:
            return json.load(file)

    def crear_documento(self, datos):
        import json

        # Leer los datos del archivo configuración
        with open('configuracion.json', 'r') as file:
            configuracion = json.load(file)

        # Extraer nombre y cédula
        nombre_terapeuta = configuracion["nombre_terapeuta"]
        cedula_profesional = configuracion["cedula_profesional"]
        
        doc = Document()
        # Configurar la fuente del documento a Arial 11
        estilo = doc.styles['Normal']
        fuente = estilo.font
        fuente.name = 'Arial'
        fuente.size = Pt(11)

        # Añadir una imagen al header del documento
        section = doc.sections[0]
        header = section.header

        # Insertar la imagen en el encabezado
        image_path_header = self.resource_path("Cabezera.png")
        header_paragraph = header.paragraphs[0]
        header_paragraph.add_run().add_picture(image_path_header, width=Inches(6.1))  # Ajusta el tamaño de la imagen

        # Añadir una imagen al footer del documento
        section = doc.sections[0]
        footer = section.footer

        # Insertar la imagen en el pie de página
        image_path_footer = self.resource_path("Fondo.png")
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.add_run().add_picture(image_path_footer, width=Inches(6.1))  # Ajusta el tamaño de la imagen
        
        # Añadir el contenido


        # Función para eliminar los bordes de una celda específica
        def remove_cell_borders(cell, borders_to_remove):
            """
            Elimina los bordes especificados de una celda.
        
            :param cell: La celda a la que se le eliminarán los bordes.
            :param borders_to_remove: Una lista de los bordes a eliminar. Ej: ['top', 'left', 'bottom', 'right']
            """
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
        
            # Crear el elemento w:tcBorders si no existe
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
        
            # Eliminar los bordes especificados
            for border in borders_to_remove:
                border_element = tcBorders.find(qn(f'w:{border}'))
                if border_element is None:
                    border_element = OxmlElement(f'w:{border}')
                    tcBorders.append(border_element)
                border_element.set(qn('w:val'), 'nil')  # Eliminar el borde

        # Añadir la tabla con 3 filas y 2 columnas
        table1 = doc.add_table(rows=3, cols=2)
        table1.style = 'Table Grid'


        # Quitar los bordes de las celdas originales antes de combinar
        remove_cell_borders(table1.cell(0, 0), ['top', 'left'])  # Borrar bordes superior e izquierdo de la celda (0, 0)
        remove_cell_borders(table1.cell(1, 0), ['left'])  # Borrar borde izquierdo de la celda (1, 0)
        remove_cell_borders(table1.cell(2, 0), ['left', 'bottom'])  # Borrar bordes izquierdo e inferior de la celda (2, 0)

        # Combinar las celdas de la primera columna (de las 3 filas)
        a = table1.cell(0, 0)  # Primera celda de la primera columna
        b = table1.cell(2, 0)  # Última celda de la primera columna
        merged_cell = a.merge(b)  # Combinar las celdas de la primera columna y asignar a merged_cell

        # Insertar la imagen en la primera columna (ocupando las 3 filas de la primera columna)
        image_path = self.resource_path("Tabla.png")
        merged_cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2.5))

        # Llenar los nombres de los datos en la segunda columna y los valores en la tercera columna

        # Formateo de la celda (0, 1)
        cell_0_1 = table1.cell(0, 1).paragraphs[0]
        run_0_1 = cell_0_1.add_run('Fecha: ')
        run_0_1.bold = True
        cell_0_1.add_run(datos["Fecha"])

        # Formateo de la celda (1, 1)
        cell_1_1 = table1.cell(1, 1).paragraphs[0]
        run_1_1 = cell_1_1.add_run('Número de Expediente: ')
        run_1_1.bold = True
        cell_1_1.add_run(datos["Número de Expediente"])

        # Formateo de la celda (2, 1)
        cell_2_1 = table1.cell(2, 1).paragraphs[0]
        run_2_1 = cell_2_1.add_run('Unidad: ')
        run_2_1.bold = True
        cell_2_1.add_run(datos["Unidad"])
     
        # Añadir una tabla con 1 fila y 1 columna
        table6 = doc.add_table(rows=1, cols=1)
        cell = table6.cell(0, 0)

        # Cambiar el color de fondo de la celda a guinda
        tcPr = cell._element.get_or_add_tcPr()  # Obtener las propiedades de la celda
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '800000')  # Código hexadecimal para guinda (color de fondo)
        tcPr.append(shading_elm)

        # Añadir texto a la celda
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

        # Añadir el texto "PLAN DE INTERVENCIÓN PSICOSOCIAL"
        run = paragraph.add_run("PLAN DE INTERVENCIÓN PSICOSOCIAL")
        run.bold = True
        run.font.size = Pt(11)  # Tamaño de la fuente
        run.font.color.rgb = RGBColor(255, 255, 255)  # Color del texto (blanco)

        # Espacio entre tablas
        doc.add_paragraph('')

        # Segunda tabla: Encabezados en la primera columna y valores en la segunda columna
        table2 = doc.add_table(rows=3, cols=2)
        table2.style = 'Table Grid'

        # Encabezados en la primera columna

        # Formateo de la celda (0, 0)
        cell = table2.cell(0, 0).paragraphs[0]
        run = cell.add_run('Nombre: ')
        run.bold = True
        cell.add_run(datos["Nombre del Paciente"])

        # Formateo de la celda (0, 1)
        cell = table2.cell(0, 1).paragraphs[0]
        run = cell.add_run('Edad: ')
        run.bold = True
        cell.add_run(datos["Edad"])

        # Formateo de la celda (1, 0)
        cell = table2.cell(1, 0).paragraphs[0]
        run = cell.add_run('Fecha de Nacimiento: ')
        run.bold = True
        cell.add_run(datos["Fecha de Nacimiento"])

        # Formateo de la celda (1, 1)
        cell = table2.cell(1, 1).paragraphs[0]
        run = cell.add_run('Diagnóstico: ')
        run.bold = True
        cell.add_run(datos["Diagnóstico"])

        # Formateo de la celda (2, 1)
        cell = table2.cell(2, 1).paragraphs[0]
        run = cell.add_run('Num. de Sesiones por semana: ')
        run.bold = True
        cell.add_run(datos["Sesiones por Semana"])

        # Formateo de la celda (2, 0)
        cell = table2.cell(2, 0).paragraphs[0]
        run = cell.add_run('Intervención en terapia: ')
        run.bold = True
        cell.add_run(datos["Área de Intervención"])
        
        # Espacio entre tablas
        doc.add_paragraph('')

        # Añadir una tabla con 1 fila y 1 columna
        table7 = doc.add_table(rows=1, cols=1)
        cell = table7.cell(0, 0)

        # Cambiar el color de fondo de la celda a guinda
        tcPr = cell._element.get_or_add_tcPr()  # Obtener las propiedades de la celda
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '800000')  # Código hexadecimal para guinda (color de fondo)
        tcPr.append(shading_elm)

        # Añadir texto a la celda
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

        # Añadir el texto "REFORZADORES"
        run = paragraph.add_run("REFORZADORES")
        run.bold = True
        run.font.size = Pt(11)  # Tamaño de la fuente
        run.font.color.rgb = RGBColor(255, 255, 255)  # Color del texto (blanco)



        # Tercera tabla: Encabezados en la primera columna y valores en la segunda columna
        table1 = doc.add_table(rows=3, cols=2)
        table1.style = 'Table Grid'
        
        # Encabezados en la primera columna
        hdr_cells1 = table1.rows[0].cells
        hdr_cells1[0].text = 'Comestibles'
        hdr_cells1[1].text = datos['Reforzadores Comestibles']

        row_cells1 = table1.rows[1].cells
        row_cells1[0].text = 'Tangibles'
        row_cells1[1].text = datos['Reforzadores Tangibles']

        row_cells2 = table1.rows[2].cells
        row_cells2[0].text = 'Sociales'
        row_cells2[1].text = datos['Reforzadores Sociales']

        # Espacio entre tablas
        doc.add_paragraph('')

        # Crear una tabla de una celda para las observaciones clínicas
        table9 = doc.add_table(rows=1, cols=1)
        table9.style = 'Table Grid'

        # Eliminar los bordes de la tabla
        def remove_borders(table):
            tbl = table._element
            tblPr = tbl.xpath('.//w:tblPr')[0]  # Buscar el elemento tblPr
            tblBorders = OxmlElement('w:tblBorders')
        
            # Crear los bordes vacíos
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')  # Configura el valor a 'none' para eliminar bordes
                tblBorders.append(border)
        
            tblPr.append(tblBorders)  # Agregar la configuración de bordes a la tabla

        # Insertar las observaciones clínicas en la tabla
        cell = table9.cell(0, 0)
        cell.text = ('El plan de intervención en terapia integral tiene como objetivo estimular habilidades cognitivas '
                 'relacionadas con los procesos mentales implicados en la ejecución y planeación de tareas o actividades, '
                 'permitiendo adaptar la conducta ante los diferentes contextos sociales. Así como el desarrollo de habilidades '
                 'relacionadas con el lenguaje, la comprensión y expresión de ideas. Varía según el perfil del alumno.')

        # Llamar la función para eliminar los bordes de la tabla
        remove_borders(table9)

        # Espacio entre tablas
        doc.add_paragraph('')
        
        # Sección de observaciones clínicas
        parrafo = doc.add_paragraph()
        run = parrafo.add_run('Observaciones Clínicas')
        run.bold = True  # Poner en negritas
        run.underline = True  # Subrayar

        # Crear una tabla de una celda para las observaciones clínicas
        table4 = doc.add_table(rows=1, cols=1)
        table4.style = 'Table Grid'

        # Insertar las observaciones clínicas en la tabla
        cell = table4.cell(0, 0)
        cell.text = datos['Observaciones Clínicas']

        # Insertar un salto de página
        doc.add_page_break()

        # Añadir una tabla con 1 fila y 1 columna
        table8 = doc.add_table(rows=1, cols=1)
        cell = table8.cell(0, 0)

        # Cambiar el color de fondo de la celda a guinda
        tcPr = cell._element.get_or_add_tcPr()  # Obtener las propiedades de la celda
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '800000')  # Código hexadecimal para guinda (color de fondo)
        tcPr.append(shading_elm)

        # Añadir texto a la celda
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto

        # Añadir el texto "Plan de Trabajo"
        run = paragraph.add_run("Plan de Trabajo")
        run.bold = True
        run.font.size = Pt(11)  # Tamaño de la fuente
        run.font.color.rgb = RGBColor(255, 255, 255)  # Color del texto (blanco)

        
        # Crear la tabla con borde continuo
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = 'Table Grid'

        # Aplicar bordes a las celdas
        for fila in tabla.rows:
            for celda in fila.cells:
                celda._element.get_or_add_tcPr().append(parse_xml(
                    r'<w:tcBorders {}>'
                    r'<w:top w:val="single" w:sz="4"/>'
                    r'<w:left w:val="single" w:sz="4"/>'
                    r'<w:bottom w:val="single" w:sz="4"/>'
                    r'<w:right w:val="single" w:sz="4"/>'
                    r'</w:tcBorders>'.format(nsdecls('w'))))

        # Añadir encabezados a la tabla
        encabezado = tabla.rows[0].cells
        encabezado[0].text = 'Nombre'
        encabezado[1].text = 'Objetivo'
        encabezado[2].text = 'Procedimiento'
        encabezado[3].text = 'Ayudas'

        # Aplicar negritas a los textos de los encabezados
        for celda in encabezado:
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.bold = True

        # Añadir los programas seleccionados
        for programa_id in datos["Programas Seleccionados"]:
            # Convertir programa_id a entero para asegurarse de que coincida con las claves del diccionario
            programa_id = int(programa_id)

            # Verificar si el programa_id existe en el diccionario
            if programa_id not in programas:
                print(f"Error: El programa_id '{programa_id}' no se encuentra en el diccionario.")
                continue

            programa = programas[programa_id]
            fila = tabla.add_row().cells
            nombre_parrafo = fila[0].paragraphs[0]
            run_nombre = nombre_parrafo.add_run(programa["Nombre"])
            run_nombre.bold = True
            fila[1].text = programa["Objetivo"]
            fila[2].text = programa["Procedimiento"]
            fila[3].text = programa["Ayudas"]


        # Añadir filas al final de la tabla con Nombre y Cédula
        fila_nombre = tabla.add_row().cells
        fila_cedula = tabla.add_row().cells

        # Añadir texto centrado y formateado para Nombre
        fila_nombre[0].merge(fila_nombre[3])  # Combinar todas las celdas de la fila
        parrafo_nombre = fila_nombre[0].paragraphs[0]
        parrafo_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto
        run_nombre = parrafo_nombre.add_run(f"Nombre: {configuracion['nombre_terapeuta']}")
        run_nombre.bold = True
        run_nombre.font.size = Pt(11)  # Tamaño de fuente

        # Añadir texto centrado y formateado para Cédula
        fila_cedula[0].merge(fila_cedula[3])  # Combinar todas las celdas de la fila
        parrafo_cedula = fila_cedula[0].paragraphs[0]
        parrafo_cedula.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto
        run_cedula = parrafo_cedula.add_run(f"Cédula: {configuracion['cedula_profesional']}")
        run_cedula.bold = True
        run_cedula.font.size = Pt(11)  # Tamaño de fuente

        # Aplicar bordes continuos a las celdas combinadas
        for fila in [fila_nombre, fila_cedula]:
            celda = fila[0]
            celda._element.get_or_add_tcPr().append(parse_xml(
                r'<w:tcBorders {}>'
                r'<w:top w:val="single" w:sz="4"/>'  # Borde superior
                r'<w:left w:val="single" w:sz="4"/>'  # Borde izquierdo
                r'<w:bottom w:val="single" w:sz="4"/>'  # Borde inferior
                r'<w:right w:val="single" w:sz="4"/>'  # Borde derecho
                r'</w:tcBorders>'.format(nsdecls('w'))
            ))

        # Guardar el documento
        doc.save(f'plan_de_trabajo_{datos["Número de Expediente"]}.docx')
